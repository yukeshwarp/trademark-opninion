# Version - 6.1  (Code Updated of Checking for Non-matching Class number == GPT 4o mini)

import os

os.environ["STREAMLIT_SERVER_WATCH_DIRS"] = "false"

from fileinput import filename
import time, os
import streamlit as st
import pandas as pd
import fitz
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re, ast
import nltk
from dotenv import load_dotenv

load_dotenv()

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

nltk.download("wordnet")
nltk.download("omw-1.4")  # Open Multilingual WordNet is also often needed


class TrademarkDetails(BaseModel):
    trademark_name: str = Field(
        description="The name of the Trademark", example="DISCOVER"
    )
    status: str = Field(description="The Status of the Trademark", example="Registered")
    serial_number: str = Field(
        description="The Serial Number of the trademark from Chronology section",
        example="87−693,628",
    )
    international_class_number: List[int] = Field(
        description="The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section",
        example=[18],
    )
    owner: str = Field(
        description="The owner of the trademark", example="WALMART STORES INC"
    )
    goods_services: str = Field(
        description="The goods/services from the document",
        example="LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES",
    )
    page_number: int = Field(
        description="The page number where the trademark details are found in the document",
        example=3,
    )
    registration_number: Union[str, None] = Field(
        description="The Registration number of the trademark from Chronology section",
        example="5,809,957",
    )
    design_phrase: str = Field(
        description="The design phrase of the trademark",
        example="THE MARK CONSISTS OF THE STYLIZED WORD 'MINI' FOLLOWED BY 'BY MOTHERHOOD.'",
        default="",
    )


# azure_endpoint = st.secrets["AZURE_ENDPOINT"]
# api_key = st.secrets["AZURE_API_KEY"]


def preprocess_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"[\u2013\u2014]", "-", text)
    return text


def is_correct_format_code1(page_text: str) -> bool:
    required_fields = ["Status:", "Goods/Services:"]  # , "Last Reported Owner:"
    return all(field in page_text for field in required_fields)


def is_correct_format_code2(page_text: str) -> bool:
    required_fields = ["Register", "Nice Classes", "Goods & Services"]
    return all(field in page_text for field in required_fields)


def extract_trademark_details_code1(
    document_chunk: str,
) -> Dict[str, Union[str, List[int]]]:
    try:
        from openai import AzureOpenAI

        # azure_endpoint = st.secrets["AZURE_OPENAI_ENDPOINT"]
        # api_key = st.secrets["AZURE_OPENAI_API_KEY"]
        azure_endpoint = os.getenv("AZURE_ENDPOINT")
        api_key = os.getenv("AZURE_API_KEY")

        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-08-01-preview",
        )

        messages = [
            {
                "role": "system",
                "content": "You are a helpful assistant for extracting Meta Data from the Trademark Document.",
            },
            {
                "role": "user",
                "content": f"""
                Extract the following details from the trademark document: trademark name, status.\n\nDocument:\n{document_chunk}
                Don't extract the same trademark details more than once; extract them only once. 
                 
                Return output only in the below mentioned format:
                Example-1 output format: 
                    Trademark Name: SLIK\n 
                    Status: PENDING\n
                Example-2 output format: 
                    Trademark Name: HUMOR US GOODS\n 
                    Status: REGISTERED\n
                Example-3 output format: 
                    Trademark Name: #WASONUO %& PIC\n 
                    Status: REGISTERED\n
                Example-4 output format: 
                    Trademark Name: AT Present, WE'VE GOT YOUR-BACK(SIDE)\n 
                    Status: PUBLISHED\n\n
                    
                Note: The trademark name length can also be 1 or 2 characters. (Example: Trademark Name: PI), (Example: Trademark Name: PII) \n""",
            },
        ]

        # Example-5 output format:
        #     Trademark Name: PI\n
        #     Status: REGISTERED\n

        # Not available in the provided document
        #  Example expected output format: Trademark Name: SLIK Status: PENDING FILED AS USE APPLICATION Serial Number: 98-602,112 International Class Number: 3 Owner: SLIK DE VENEZUELA C.A. VENEZUELA CORPORATION Goods & Services: Cosmetics; hair gel; hair wax; hair styling gel; non-medicated cosmetics Filed Date: JUN 14, 2024 Registration Number: Not available in the provided document.
        #  Example expected output: Trademark Name: #WASONOFILTER Status: REGISTERED Serial Number: 88-404,432 International Class Number: 21 Owner: LAHC US 1 LLC DELAWARE LIMITED LIABILITY COMPANY Goods & Services:  Containers for household use, coffee mugs, and wine glasses Filed Date: APR 26, 2019 Registration Number: 5,963,355"""}

        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=messages,
            temperature=0,
            max_tokens=300,
        )
        extracted_text = response.choices[0].message.content

        # if extracted_text and extracted_text != "[]":
        # st.write(extracted_text)

        details = {}
        for line in extracted_text.split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                details[key.strip().lower().replace(" ", "_")] = value.strip()

        # st.warning(details)
        return details

    except Exception as e:
        print(f"An error occurred: {e}")


def extract_serial_number(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    combined_texts = ""
    for i in range(start_page, min(start_page + 13, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Serial Number:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Chronology:.*?Serial Number:\s*([\d,-−]+)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        registration_number = match.group(1).strip()
        return registration_number
    return "No serial number presented in document"


def extract_ownership(
    document: str, start_page: int, proposed_name: str, pdf_document: fitz.Document
) -> str:
    """Extract the ownership from the document"""
    combined_texts = ""
    for i in range(start_page, min(start_page + 13, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Last Reported Owner:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Last Reported Owner:\s*(.*?)\n\s*(.*?)\n"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        owner_name = match.group(1).strip()
        owner_type = match.group(2).strip()
        if owner_type == proposed_name:
            return f"{owner_name}"
        else:
            return f"{owner_name} {owner_type}"
    return "Not available in the provided document."


def extract_registration_number(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    """Extract the registration number from the document"""
    combined_texts = ""
    for i in range(start_page, min(start_page + 8, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Registration Number:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Last ReportedOwner:.*?Registration Number:\s*([\d,]+)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        registration_number = match.group(1).strip()
        return registration_number
    return "NA"


def extract_trademark_details_code2(page_text: str) -> Dict[str, Union[str, List[int]]]:
    details = {}

    trademark_name_match = re.search(
        r"\d+\s*/\s*\d+\s*\n\s*\n\s*([A-Za-z0-9'&!,\-. ]+)\s*\n", page_text
    )
    if trademark_name_match:
        details["trademark_name"] = trademark_name_match.group(1).strip()
    else:
        trademark_name_match = re.search(
            r"(?<=\n)([A-Za-z0-9'&!,\-. ]+)(?=\n)", page_text
        )
        details["trademark_name"] = (
            trademark_name_match.group(1).strip() if trademark_name_match else ""
        )

    status_match = re.search(
        r"Status\s*(?:\n|:\s*)([A-Za-z]+)", page_text, re.IGNORECASE
    )
    details["status"] = status_match.group(1).strip() if status_match else ""

    owner_match = re.search(r"Holder\s*(?:\n|:\s*)(.*)", page_text, re.IGNORECASE)
    if owner_match:
        details["owner"] = owner_match.group(1).strip()
    else:
        owner_match = re.search(r"Owner\s*(?:\n|:\s*)(.*)", page_text, re.IGNORECASE)
        details["owner"] = owner_match.group(1).strip() if owner_match else ""

    nice_classes_match = re.search(
        r"Nice Classes\s*[\s:]*\n((?:\d+(?:,\s*\d+)*)\b)", page_text, re.IGNORECASE
    )
    if nice_classes_match:
        nice_classes_text = nice_classes_match.group(1)
        nice_classes = [int(cls.strip()) for cls in nice_classes_text.split(",")]
        details["international_class_number"] = nice_classes
    else:
        details["international_class_number"] = []

    serial_number_match = re.search(r"Application#\s*(.*)", page_text, re.IGNORECASE)
    details["serial_number"] = (
        serial_number_match.group(1).strip() if serial_number_match else ""
    )

    goods_services_match = re.search(
        r"Goods & Services\s*(.*?)(?=\s*G&S translation|$)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    details["goods_services"] = (
        goods_services_match.group(1).strip() if goods_services_match else ""
    )

    registration_number_match = re.search(
        r"Registration#\s*(.*)", page_text, re.IGNORECASE
    )
    details["registration_number"] = (
        registration_number_match.group(1).strip() if registration_number_match else ""
    )

    # Description
    design_phrase = re.search(
        r"Description\s*(.*?)(?=\s*Applicant|Owner|Holder|$)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    details["design_phrase"] = (
        design_phrase.group(1).strip()
        if design_phrase
        else "No Design phrase presented in document"
    )

    return details


def read_pdf(file_path: str, exclude_header_footer: bool = True) -> str:
    document_text = ""
    with fitz.open(file_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            if exclude_header_footer:
                rect = page.rect
                x0 = rect.x0
                y0 = rect.y0 + rect.height * 0.1
                x1 = rect.x1
                y1 = rect.y1 - rect.height * 0.1
                page_text = page.get_text("text", clip=(x0, y0, x1, y1))
            else:
                page_text = page.get_text()
            document_text += page_text
    return document_text


def split_text(text: str, max_tokens: int = 1500) -> List[str]:
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split("\n"):
        line_length = len(line.split())
        if current_length + line_length > max_tokens:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]
            current_length = line_length
        else:
            current_chunk.append(line)
            current_length += line_length

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def parse_international_class_numbers(class_numbers: str) -> List[int]:
    numbers = class_numbers.split(",")
    return [int(num.strip()) for num in numbers if num.strip().isdigit()]


def extract_international_class_numbers_and_goods_services(
    document: str, start_page: int, pdf_document: fitz.Document
) -> Dict[str, Union[List[int], str]]:
    """Extract the International Class Numbers and Goods/Services from the document over a range of pages"""
    class_numbers = []
    goods_services = []
    combined_text = ""

    for i in range(start_page, min(start_page + 10, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_text += page_text
        if "Last Reported Owner:" in page_text:
            break

    pattern = r"International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|Disclaimers:|\Z)"
    matches = re.findall(pattern, combined_text, re.DOTALL)
    for match in matches:
        class_number = int(match[0])
        class_numbers.append(class_number)
        goods_services.append(f"Class {class_number}: {match[1].strip()}")

    if "sexual" in goods_services or "sex" in goods_services:
        goods_services = replace_disallowed_words(goods_services)

    return {
        "international_class_numbers": class_numbers,
        "goods_services": "\n".join(goods_services),
    }


def extract_design_phrase(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    """Extract the design phrase from the document"""
    combined_texts = ""
    for i in range(start_page, min(start_page + 10, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Design Phrase:" in page_text or "Filing Correspondent:" in page_text:
            break

    pattern = r"Design Phrase:\s*(.*?)(?=Other U\.S\. Registrations:|Filing Correspondent:|Group:|USPTO Page:|$)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        design_phrase = match.group(1).strip()
        # Remove any newline characters within the design phrase
        design_phrase = " ".join(design_phrase.split())
        return design_phrase
    return "No Design phrase presented in document"


def parse_trademark_details(
    document_path: str,
) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()

            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = (
                    extract_international_class_numbers_and_goods_services(
                        page_text, page_num, pdf_document
                    )
                )
                registration_number = extract_registration_number(
                    page_text, page_num, pdf_document
                )
                serial_number = extract_serial_number(page_text, page_num, pdf_document)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                ownership_details = extract_ownership(
                    page_text, page_num, proposed_name, pdf_document
                )

                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    extracted_data["owner"] = ownership_details
                    extracted_data["serial_number"] = serial_number
                    extracted_data["registration_number"] = registration_number
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "")
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[
                                0
                            ].strip()
                        if re.match(r"^US-\d+", trademark_name):
                            trademark_name = re.sub(
                                r"^US-\d+\s*", "", trademark_name
                            ).strip()
                        status = data.get("status", "").split(",")[0].strip()
                        serial_number = data.get("serial_number", "")
                        owner = data.get("owner", "")
                        international_class_number = data.get(
                            "international_class_numbers", []
                        )
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get(
                            "registration_number",
                            "No registration number presented in document",
                        )
                        design_phrase = data.get(
                            "design_phrase", "No Design phrase presented in document"
                        )

                        # If crucial fields are missing, attempt to re-extract the values

                        # if not trademark_name or not owner or not status or not international_class_number:
                        #     preprocessed_chunk = preprocess_text(data.get("raw_text", ""))
                        #     extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                        #     trademark_name = extracted_data.get("trademark_name", trademark_name).split(',')[0].strip()
                        #     if "Global Filings" in trademark_name:
                        #         trademark_name = trademark_name.split("Global Filings")[0].strip()
                        #     owner = extracted_data.get("owner", owner).split(',')[0].strip()
                        #     status = extracted_data.get("status", status).split(',')[0].strip()
                        #     international_class_number = parse_international_class_numbers(extracted_data.get("international_class_number", "")) or international_class_number
                        #     registration_number = extracted_data.get("registration_number", registration_number).split(',')[0].strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase,
                        )
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number": trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase,
                        }
                        print(trademark_info)
                        print(
                            "_____________________________________________________________________________________________________________________________"
                        )
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

            else:
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                st.info("Code 2")
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get(
                                "international_class_number", []
                            ),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", ""),
                        )
                        if (
                            trademark_details.trademark_name != ""
                            and trademark_details.owner != ""
                            and trademark_details.status != ""
                            and trademark_details.goods_services != ""
                        ):
                            trademark_info = {
                                "trademark_name": trademark_details.trademark_name,
                                "owner": trademark_details.owner,
                                "status": trademark_details.status,
                                "serial_number": trademark_details.serial_number,
                                "international_class_number": trademark_details.international_class_number,
                                "goods_services": trademark_details.goods_services,
                                "page_number": trademark_details.page_number,
                                "registration_number": trademark_details.registration_number,
                                "design_phrase": trademark_details.design_phrase,
                            }

                            trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list


from typing import List, Dict, Union
from sentence_transformers import SentenceTransformer, util
from fuzzywuzzy import fuzz
import torch

# Load the semantic similarity model with device specified during initialization
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
semantic_model = SentenceTransformer(
    "sentence-transformers/all-MiniLM-L6-v2", device=device
)


# Rest of your code remains the same
def compare_trademarks(
    existing_trademark: Dict[str, Union[str, List[int]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> Dict[str, Union[str, int]]:
    # Convert proposed classes to a list of integers
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]

    # Helper function for semantic equivalence
    def is_semantically_equivalent(name1, name2, threshold=0.80):
        embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        return similarity_score >= threshold, similarity_score + 30

    # Helper function for phonetic equivalence
    def is_phonetically_equivalent(name1, name2, threshold=80):
        score = fuzz.ratio(name1.lower(), name2.lower())
        return score >= threshold, score

    # Helper function for phonetically equivalent words
    def first_words_phonetically_equivalent(existing_name, proposed_name, threshold=80):
        existing_words = existing_name.lower().split()
        proposed_words = proposed_name.lower().split()
        if len(existing_words) < 2 or len(proposed_words) < 2:
            return False
        return (
            fuzz.ratio(" ".join(existing_words[:2]), " ".join(proposed_words[:2]))
            >= threshold
        )

    # Condition 1A: Exact character-for-character match
    condition_1A_satisfied = (
        existing_trademark["trademark_name"].strip().lower()
        == proposed_name.strip().lower()
    )

    # Condition 1B: Semantically equivalent
    condition_1B_satisfied, condition_1B_score = is_semantically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1C: Phonetically equivalent
    condition_1C_satisfied, condition_1C_score = is_phonetically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1D: First two or more words are phonetically equivalent
    condition_1D_satisfied = first_words_phonetically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1E: Proposed name is the first word of the existing trademark
    condition_1E_satisfied = (
        existing_trademark["trademark_name"].lower().startswith(proposed_name.lower())
    )

    # Check if any Condition 1 is satisfied
    condition_1_satisfied = any(
        [
            condition_1A_satisfied,
            condition_1B_satisfied,
            condition_1C_satisfied,
            condition_1D_satisfied,
            condition_1E_satisfied,
        ]
    )

    # Condition 2: Overlap in International Class Numbers
    condition_2_satisfied = bool(
        set(existing_trademark["international_class_number"]) & set(proposed_classes)
    )

    import re
    from nltk.stem import WordNetLemmatizer

    def normalize_text(text):

        # Replace special hyphen-like characters with a standard hyphen
        text = re.sub(r"[−–—]", "-", text)
        # Remove punctuation except hyphens and spaces
        text = re.sub(r"[^\w\s-]", " ", text)
        # Convert to lowercase
        text = text.lower()
        text = re.sub(r"\b\d+\b", "", text)
        text = re.sub(r"\bclass\b", "", text)
        text = re.sub(r"\bcare\b", "", text)
        text = re.sub(r"\bin\b", "", text)
        text = re.sub(r"\band\b", "", text)
        text = re.sub(r"\bthe\b", "", text)
        text = re.sub(r"\bfor\b", "", text)
        text = re.sub(r"\bwith\b", "", text)
        text = re.sub(r"\bfrom\b", "", text)
        text = re.sub(r"\bto\b", "", text)
        text = re.sub(r"\bunder\b", "", text)
        text = re.sub(r"\busing\b", "", text)
        text = re.sub(r"\bof\b", "", text)
        text = re.sub(r"\bno\b", "", text)
        text = re.sub(r"\binclude\b", "", text)
        text = re.sub(r"\bex\b", "", text)
        text = re.sub(r"\bexample\b", "", text)
        text = re.sub(r"\bclasses\b", "", text)
        text = re.sub(r"\bsearch\b", "", text)
        text = re.sub(r"\bscope\b", "", text)
        text = re.sub(r"\bshower\b", "", text)
        text = re.sub(r"\bproducts\b", "", text)
        text = re.sub(r"\bshampoos\b", "hair", text)

        # Standardize whitespace
        return " ".join(text.split())

    # Condition 3: Target market and goods/services overlap
    def target_market_and_goods_overlap(existing_gs, proposed_gs, threshold=0.65):

        existing_normalized = normalize_text(existing_gs)
        proposed_normalized = normalize_text(proposed_gs)

        embeddings1 = semantic_model.encode(existing_normalized, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(proposed_normalized, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        # st.write("Semantic Similarity Score:", similarity_score)
        if similarity_score >= threshold:
            return True

        # Split into words and lemmatize
        lemmatizer = WordNetLemmatizer()
        existing_words = {
            lemmatizer.lemmatize(word) for word in existing_normalized.split()
        }
        proposed_words = {
            lemmatizer.lemmatize(word) for word in proposed_normalized.split()
        }

        # Check for common words
        common_words = existing_words.intersection(proposed_words)
        # st.write("Common Words:", existing_gs , common_words)
        return bool(common_words)

    condition_3_satisfied = target_market_and_goods_overlap(
        existing_trademark["goods_services"], proposed_goods_services
    )

    # condition_1A_satisfieds = is_exact_match(existing_trademark['trademark_name'].strip().lower(), proposed_name.strip().lower())
    # st.write(f"Exact Match: {condition_1A_satisfieds}")

    # condition_1B_satisfieds = is_semantically_equivalents(existing_trademark['trademark_name'].strip().lower(), proposed_name.strip().lower())
    # st.write(f"Semantically equivalents : {condition_1B_satisfieds}")

    # condition_1C_satisfieds = is_phonetically_equivalents(existing_trademark['trademark_name'], proposed_name)
    # st.write(f"Phonetically equivalents : {condition_1C_satisfieds}")

    # condition_3_satisfieds = target_market_and_goods_overlaps(existing_trademark['goods_services'], proposed_goods_services)
    # st.write(f"Goods and services match's : {condition_3_satisfieds}")

    # Clean and standardize the trademark status
    status = existing_trademark["status"].strip().lower()

    # Check for 'Cancelled' or 'Abandoned' status
    if any(keyword in status for keyword in ["cancelled", "abandoned", "expired"]):
        conflict_grade = "Low"
        reasoning = "The existing trademark status is 'Cancelled' or 'Abandoned.'"
    else:
        points = sum(
            [
                condition_1_satisfied,  # 1 point if any Condition 1 is satisfied
                condition_2_satisfied,  # 1 point if Condition 2 is satisfied
                condition_3_satisfied,  # 1 point if Condition 3 is satisfied
            ]
        )

        # Determine conflict grade based on points
        if points == 3:
            conflict_grade = "High"
        elif points == 2:
            conflict_grade = "Moderate"
        elif points == 1:
            conflict_grade = "Low"
        else:
            conflict_grade = "None"

        if condition_1_satisfied:
            condition_1_details = []
            if condition_1A_satisfied:
                condition_1_details.append("Exact character-for-character match")
            if condition_1B_satisfied:
                condition_1_details.append("Semantically equivalent")
            if condition_1C_satisfied:
                condition_1_details.append("Phonetically equivalent")
            if condition_1D_satisfied:
                condition_1_details.append(
                    "First two or more words are phonetically equivalent"
                )
            if condition_1E_satisfied:
                condition_1_details.append(
                    "Proposed name is the first word of the existing trademark"
                )

        # Generate detailed reasoning for Condition 1
        if condition_1_satisfied:
            condition_1_reasoning = (
                f"Condition 1: Satisfied - {', '.join(condition_1_details)}."
            )
        else:
            condition_1_reasoning = "Condition 1: Not Satisfied."

        # Reasoning
        reasoning = (
            f"{condition_1_reasoning} \n"
            f"Condition 2: {'Satisfied' if condition_2_satisfied else 'Not Satisfied'} - Overlap in class numbers.\n"
            f"Condition 3: {'Satisfied' if condition_3_satisfied else 'Not Satisfied'} - Overlap in goods/services and target market."
        )

    if existing_trademark["design_phrase"] == "No Design phrase presented in document":
        design_label = "Word"
    else:
        design_label = "Design"

    if condition_1_satisfied and condition_2_satisfied and condition_3_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": "   ✔️",
            "Class": "   ✔️",
            "Goods/Services": "   ✔️",
            "Direct Hit": " ",
            "semantic_score": condition_1B_score,
            "phonetic_score": condition_1C_score,
        }

    elif condition_1_satisfied and condition_2_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": "   ✔️",
            "Class": "   ✔️",
            "Goods/Services": "  ",
            "Direct Hit": " ",
            "semantic_score": condition_1B_score,
            "phonetic_score": condition_1C_score,
        }

    elif condition_2_satisfied and condition_3_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": " ",
            "Class": "   ✔️",
            "Goods/Services": "   ✔️",
            "Direct Hit": " ",
            "semantic_score": condition_1B_score,
            "phonetic_score": condition_1C_score,
        }

    else:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": " ",
            "Class": "   ✔️",
            "Goods/Services": " ",
            "Direct Hit": " ",
            "semantic_score": condition_1B_score,
            "phonetic_score": condition_1C_score,
        }


def replace_disallowed_words(text):
    disallowed_words = {
        "sexual": "xxxxxx",
        "sex": "xxx",
    }
    for word, replacement in disallowed_words.items():
        text = text.replace(word, replacement)
    # Ensure single paragraph output
    text = " ".join(text.split())
    return text


def assess_conflict(
    existing_trademark: List[Dict[str, Union[str, List[int]]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> List[Dict[str, int]]:

    import phonetics
    from sentence_transformers import util
    from rapidfuzz import fuzz

    def normalize_text_name(text):
        """Normalize text by converting to lowercase, removing special characters, and standardizing whitespace."""
        # Remove punctuation except hyphens and spaces
        # text = re.sub(r"[^\w\s-']", "", text)
        # Convert to lowercase
        text = re.sub(r"'", " ", text)
        text = text.lower()
        # Standardize whitespace
        return " ".join(text.split())

    # Clean and standardize the trademark status
    status = existing_trademark["status"].strip().lower()
    semantic_similarity = 0
    string_similarity = 0
    # Check for 'Cancelled' or 'Abandoned' status
    if any(keyword in status for keyword in ["cancelled", "abandoned", "expired"]):
        conflict_grade = "Low"
        reasoning = "The existing trademark status is 'Cancelled' or 'Abandoned.'"
    else:

        existing_trademark_name = normalize_text_name(
            existing_trademark["trademark_name"]
        )
        proposed_name = normalize_text_name(proposed_name)

        # Phonetic Comparison
        existing_phonetic = phonetics.metaphone(existing_trademark_name)
        proposed_phonetic = phonetics.metaphone(proposed_name)
        phonetic_match = existing_phonetic == proposed_phonetic

        # Semantic Similarity
        existing_embedding = semantic_model.encode(
            existing_trademark_name, convert_to_tensor=True
        )
        proposed_embedding = semantic_model.encode(
            proposed_name, convert_to_tensor=True
        )
        semantic_similarity = util.cos_sim(
            existing_embedding, proposed_embedding
        ).item()

        # String Similarity
        string_similarity = fuzz.ratio(existing_trademark_name, proposed_name)

        def is_substring_match(name1, name2):
            return name1.lower() in name2.lower() or name2.lower() in name1.lower()

        substring_match = is_substring_match(existing_trademark_name, proposed_name)

        def has_shared_word(name1, name2):
            words1 = set(name1.lower().split())
            words2 = set(name2.lower().split())
            return not words1.isdisjoint(words2)

        shared_word = has_shared_word(existing_trademark_name, proposed_name)

        from fuzzywuzzy import fuzz

        def is_phonetic_partial_match(name1, name2, threshold=55):
            return fuzz.partial_ratio(name1.lower(), name2.lower()) >= threshold

        phonetic_partial_match = is_phonetic_partial_match(
            existing_trademark_name, proposed_name
        )

        # st.write(f"Shared word : {existing_trademark_name} : {shared_word}")
        # st.write(f"Phonetic partial match : {existing_trademark_name} : {phonetic_partial_match}")
        # st.write(f"Substring match : {existing_trademark_name} : {substring_match}")

        # Decision Logic
        if (
            phonetic_match
            or substring_match
            or shared_word
            or semantic_similarity >= 0.5
            or string_similarity >= 55
            or phonetic_partial_match >= 55
        ):
            conflict_grade = "Name-Match"
        else:
            conflict_grade = "Low"

        semantic_similarity = semantic_similarity * 100

        # Reasoning
        reasoning = (
            f"Condition 1: {'Satisfied' if phonetic_match else 'Not Satisfied'} - Phonetic match found.\n"
            f"Condition 2: {'Satisfied' if substring_match else 'Not Satisfied'} - Substring match found.\n"
            f"Condition 3: {'Satisfied' if shared_word else 'Not Satisfied'} - Substring match found.\n"
            f"Condition 4: {'Satisfied' if phonetic_partial_match >= 55 else 'Not Satisfied'} - String similarity is ({round(phonetic_partial_match)}%).\n"
            f"Condition 5: {'Satisfied' if semantic_similarity >= 50 else 'Not Satisfied'} - Semantic similarity is ({round(semantic_similarity)}%).\n"
            f"Condition 6: {'Satisfied' if string_similarity >= 55 else 'Not Satisfied'} - String similarity is ({round(string_similarity)}%).\n"
        )

    if existing_trademark["design_phrase"] == "No Design phrase presented in document":
        design_label = "Word"
    else:
        design_label = "Design"

    return {
        "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
        "Trademark name": existing_trademark["trademark_name"],
        "Trademark Status": existing_trademark["status"],
        "Trademark Owner": existing_trademark["owner"],
        "Trademark class Number": existing_trademark["international_class_number"],
        "Trademark serial number": existing_trademark["serial_number"],
        "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
        "Trademark registration number": existing_trademark["registration_number"],
        "Trademark design phrase": existing_trademark["design_phrase"],
        "Word/Design": design_label,
        "conflict_grade": conflict_grade,
        "reasoning": reasoning,
        "Mark": " ",
        "Class": " ",
        "Goods/Services": " ",
        "Direct Hit": "   ✔️",
        "semantic_score": semantic_similarity,
        "phonetic_score": string_similarity,
    }


import os
import json
from openai import AzureOpenAI


# Function to compare trademarks
def compare_trademarks2(
    existing_trademark: List[Dict[str, Union[str, List[int]]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> List[Dict[str, Union[str, int]]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]

    # Prepare the messages for the Azure OpenAI API
    messages = [
        {
            "role": "system",
            "content": """  
            You are a trademark attorney tasked with determining a conflict grade based on the given conditions.  
            
            **Additional Instructions:**  
            
            - Consider if the proposed trademark name appears anywhere within the existing trademark name, or if significant parts of the existing trademark name appear in the proposed name.  
            - Evaluate shared words between trademarks, regardless of their position.  
            - Assess phonetic similarities, including partial matches.  
            - Consider the overall impression created by the trademarks, including similarities in appearance, sound, and meaning.  
            
            Follow the conflict grading criteria as previously outlined, assigning "Name-Match" or "Low" based on your analysis.  
            """,
        },
        {
            "role": "user",
            "content": f"""  
            Evaluate the potential conflict between the following existing trademarks and the proposed trademark.  
            
            **Proposed Trademark:**  
            - Name: "{proposed_name}"  
            
            **Existing Trademarks:**  
            - Name: "{existing_trademark['trademark_name']}"  
            - Status: "{existing_trademark['status']}"
            
            **Instructions:**  
            1. Review the proposed and existing trademark data.  
            2. Determine if the trademarks are likely to cause confusion based on the Trademark name such as Phonetic match, Semantic similarity and String similarity.  
            3. Return the output with Conflict Grade only as 'Name-Match' or 'Low', based on the reasoning. 
            4. Provide reasoning for each Conflict Grade.
            5. Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered as Conflict Grade: Low.  
            
            **Output Format:**  
                Existing Name: Name of the existing trademark.
                Reasoning: Reasoning for the conflict grade.
                Conflict Grade: Name-Match
        """,
        },
    ]

    # Initialize the Azure OpenAI client
    # azure_endpoint = st.secrets["AZURE_OPENAI_ENDPOINT"]
    # api_key = st.secrets["AZURE_OPENAI_API_KEY"]
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    if not azure_endpoint or not api_key:
        raise ValueError(
            "Azure endpoint or API key is not set in environment variables."
        )

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint, api_key=api_key, api_version="2024-10-01-preview"
    )

    # Call Azure OpenAI to get the response
    try:
        response_reasoning = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0,
            max_tokens=500,
            top_p=1,
        )

        # Extract the content from the response
        reasoning_content = response_reasoning.choices[0].message.content
        conflict_grade = reasoning_content.split("Conflict Grade:", 1)[1].strip()
        st.write(reasoning_content)

        return conflict_grade

    except Exception as e:
        print(f"Error while calling Azure OpenAI API: {e}")
        return []


def extract_proposed_trademark_details(
    file_path: str,
) -> Dict[str, Union[str, List[int]]]:
    """Extract proposed trademark details from the given input format"""
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            if "Mark Searched:" not in page_text:
                page = pdf_document.load_page(1)
                page_text = preprocess_text(page.get_text())

    name_match = re.search(
        r"Mark Searched:\s*(.*?)(?=\s*Client Name:)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    if "Goods/Services:" in page_text:
        goods_services_match = re.search(
            r"Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)",
            page_text,
            re.IGNORECASE | re.DOTALL,
        )
    else:
        goods_services_match = re.search(
            r"Goods and Services:\s*(.*?)(?=\s*Order Info)",
            page_text,
            re.IGNORECASE | re.DOTALL,
        )

    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(
            1
        ).strip()

    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers

    return proposed_details


def find_class_numbers(goods_services: str) -> List[int]:
    """Use LLM to find the international class numbers based on goods & services"""
    # Initialize AzureChatOpenAI

    from openai import AzureOpenAI

    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant for finding the International class number of provided Goods & Services.",
        },
        {
            "role": "user",
            "content": "The goods/services are: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 03"},
        {
            "role": "user",
            "content": "The goods/services are: LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 18,35"},
        {
            "role": "user",
            "content": "The goods/services are: CLASS 3: ANTIPERSPIRANTS AND DEODORANTS. (PLEASE INCLUDE CLASSES 5 AND 35 IN THE SEARCH SCOPE). Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 03,05,35"},
        {
            "role": "user",
            "content": "The goods/services are: VITAMIN AND MINERAL SUPPLEMENTS. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 05"},
        {
            "role": "user",
            "content": f"The goods/services are: {goods_services}. Find the international class numbers.",
        },
    ]
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=messages,
        temperature=0.5,
        max_tokens=150,
    )

    class_numbers_str = response.choices[0].message.content

    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(
        r"(?<!\d)\d{2}(?!\d)", class_numbers_str
    )  # Look for two-digit numbers
    class_numbers = ",".join(
        set(class_numbers)
    )  # Convert to set to remove duplicates, then join into a single string

    return class_numbers


def extract_proposed_trademark_details2(
    file_path: str,
) -> Dict[str, Union[str, List[int]]]:
    """Extract proposed trademark details from the first page of the document"""
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())

            name_match = re.search(r"Name:\s*(.*?)(?=\s*Nice Classes:)", page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(
                    1
                ).strip()

            nice_classes_match = re.search(
                r"Nice Classes:\s*(\d+(?:,\s*\d+)*)", page_text
            )
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = (
                    nice_classes_match.group(1).strip()
                )

            goods_services_match = re.search(
                r"Goods & Services:\s*(.*?)(?=\s*Registers|$)",
                page_text,
                re.IGNORECASE | re.DOTALL,
            )
            if goods_services_match:
                proposed_details["proposed_goods_services"] = (
                    goods_services_match.group(1).strip()
                )

    return proposed_details


def list_conversion(proposed_class: str) -> List[int]:

    from openai import AzureOpenAI

    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant for converting the class number string into python list of numbers.\n Respond only with python list. Example : [18,35]",
        },
        {
            "role": "user",
            "content": "The class number are: 15,89. convert the string into python list of numbers.",
        },
        {"role": "assistant", "content": "[15,89]"},
        {
            "role": "user",
            "content": f"The class number are: {proposed_class}. convert the string into python list of numbers.",
        },
    ]
    # messages = [
    # {
    #     "role": "system",
    #     "content": "You are a helpful assistant that converts strings of class numbers into Python lists of integers."
    # },
    # {
    # "role": "user",
    # "content": f"""
    #     Convert the following string of class numbers into a Python list of integers.

    #     **Instructions:**

    #     - The input is a string of numbers separated by commas (e.g., `15,89`).
    #     - **Respond only** with a Python list of integers (e.g., `[15, 89]`).
    #     - Do not include any additional text or commentary.
    #     - Ensure the numbers are integers, not strings.

    #     **Example:**

    #     - Input: "15,89"
    #     - Response: [15, 89]

    #     **Input:**

    #     "{proposed_class}"
    #     """
    # }
    # ]

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=messages,
        temperature=0,
        max_tokens=150,
    )

    lst_class = response.choices[0].message.content
    class_value = ast.literal_eval(lst_class)

    return class_value


# TAMIL CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

import os
from openai import AzureOpenAI
import json
import re


def get_azure_client():
    """Initialize and return the Azure OpenAI client."""
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    # azure_endpoint = st.secrets["AZURE_ENDPOINT"]
    # api_key = st.secrets["AZURE_API_KEY"]

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )
    return client


def parse_conflicts_array(conflicts_array):
    """
    Parse the string-based conflicts array into a structured format.

    Args:
        conflicts_array: List of strings containing trademark conflict information

    Returns:
        List of dictionaries containing structured trademark information
    """
    print("==PARSING CONFLICTS ARRAY==")
    structured_conflicts = []
    current_conflict = {}

    for line in conflicts_array:
        print("==LINE==", line)
        line = line.strip()
        if not line:  # Skip empty lines
            continue

        if line.startswith("Trademark Name :"):
            if current_conflict:  # Save previous conflict if exists
                structured_conflicts.append(current_conflict)
            current_conflict = {
                "trademark_name": line.replace("Trademark Name :", "").strip()
            }
        elif line.startswith("Trademark Status :"):
            current_conflict["status"] = line.replace("Trademark Status :", "").strip()
        elif line.startswith("Trademark Owner :"):
            current_conflict["owner"] = line.replace("Trademark Owner :", "").strip()
        elif line.startswith("Trademark Class Number :"):
            class_str = line.replace("Trademark Class Number : ", "").strip()
            try:
                current_conflict["class_number"] = class_str.replace("[", "").replace(
                    "]", ""
                )
            except:
                current_conflict["class_number"] = class_str
        elif line.startswith("Trademark serial number :"):
            current_conflict["serial_number"] = line.replace(
                "Trademark serial number :", ""
            ).strip()
        elif line.startswith("Trademark registration number :"):
            current_conflict["registration_number"] = line.replace(
                "Trademark registration number :", ""
            ).strip()
        elif line.startswith("Trademark Design phrase :"):
            current_conflict["design_phrase"] = line.replace(
                "Trademark Design phrase :", ""
            ).strip()
        elif line.startswith("Condition"):
            if "reasoning" not in current_conflict:
                current_conflict["reasoning"] = []
            current_conflict["reasoning"].append(line)

    # Add the last conflict if exists
    if current_conflict:
        structured_conflicts.append(current_conflict)

    print("COMPLETED PARSING CONFLICTS ARRAY")

    return structured_conflicts


def validate_trademark_relevance(
    conflicts_array, proposed_goods_services, proposed_name
):
    """
    Pre-filter trademarks using ML-based similarity scores and goods/services relevance.
    Returns high confidence matches (>0.85) and borderline matches (0.75-0.85).

    Args:
        conflicts_array: List of trademark conflicts (can be string or list)
        proposed_goods_services: Goods/services of the proposed trademark
        proposed_name: Name of the proposed trademark

    Returns:
        dict: Contains high_confidence_matches, borderline_matches, and excluded_count
    """
    # Parse conflicts_array if it's a string
    print("==VALIDATING TRADEMARK RELEVANCE==")
    if isinstance(conflicts_array, str):
        print("==ISINSTANCE BLOCK==")
        try:
            print("==TRY BLOCK==")
            conflicts = json.loads(conflicts_array)
        except json.JSONDecodeError:
            print("==EXCEPT BLOCK==")
            conflicts = parse_conflicts_array(conflicts_array.split("\n"))
    else:
        print("==ELSE BLOCK==")
        conflicts = parse_conflicts_array(conflicts_array)

    # Initialize lists for different match categories
    high_confidence_matches = []
    borderline_matches = []
    excluded_count = 0

    def is_similar_goods_services(existing_goods_services, proposed_goods_services):
        """Check if goods/services are similar enough to warrant analysis."""
        # Convert to lowercase for comparison
        existing = existing_goods_services.lower()
        proposed = proposed_goods_services.lower()

        # Check for exact match
        if existing == proposed:
            return True

        # Check for substring match
        if existing in proposed or proposed in existing:
            return True

        # Remove common stop words and split into words
        stop_words = {
            "the",
            "and",
            "or",
            "for",
            "in",
            "on",
            "at",
            "to",
            "of",
            "with",
            "by",
        }
        existing_words = set(
            word for word in existing.split() if word not in stop_words
        )
        proposed_words = set(
            word for word in proposed.split() if word not in stop_words
        )

        # Check for significant word overlap
        overlap = existing_words.intersection(proposed_words)
        if len(overlap) >= min(len(existing_words), len(proposed_words)) * 0.5:
            return True

        return False

    # Process each conflict
    for conflict in conflicts:
        # First check goods/services similarity
        if not is_similar_goods_services(
            conflict.get("goods_services", ""), proposed_goods_services
        ):
            excluded_count += 1
            continue
        print("==GOODS SERVICES SIMILARITY CHECK PASSED==\n")
        # Calculate semantic and phonetic scores
        semantic_score, semantic_match = ml_semantic_match(
            proposed_name, conflict["trademark_name"]
        )
        print("==SEMANTIC MATCH CHECK PASSED==")
        phonetic_score, phonetic_match = ml_phonetic_match(
            proposed_name, conflict["trademark_name"]
        )
        print("==PHONETIC MATCH CHECK PASSED==\n")

        # Categorize based on both semantic and phonetic scores
        if semantic_score > 0.80 and phonetic_score > 0.85:
            high_confidence_matches.append(conflict)
        elif (semantic_score > 0.70 and semantic_score <= 0.80) or (
            phonetic_score > 0.75 and phonetic_score <= 0.85
        ):
            borderline_matches.append(conflict)
        else:
            excluded_count += 1

    # Print filtering results
    print(f"\nFiltering Results:")
    print(f"High Confidence Matches: {len(high_confidence_matches)}")
    print(f"Borderline Matches: {len(borderline_matches)}")
    print(f"Excluded Conflicts: {excluded_count}")

    return {
        "high_confidence_matches": high_confidence_matches,
        "borderline_matches": borderline_matches,
        "excluded_count": excluded_count,
    }


def filter_by_gpt_response(conflicts, gpt_json):
    """
    Removes trademarks that GPT flagged as lacking goods/services overlap.

    Args:
        conflicts: Original list of trademark conflicts
        gpt_json: JSON object from GPT with 'results' key

    Returns:
        Filtered list of conflicts that GPT identified as overlapping
    """
    # Parse the GPT response if it's a string
    if isinstance(gpt_json, str):
        try:
            gpt_json = json.loads(gpt_json)
        except json.JSONDecodeError:
            # If JSON is invalid, keep original conflicts
            return conflicts

    gpt_results = gpt_json.get("results", [])

    # Build a set of marks with overlap for quick membership checking
    overlapping_marks = {
        result["mark"] for result in gpt_results if result.get("overlap") is True
    }

    # Retain conflicts only if they appear in overlapping_marks
    filtered_conflicts = [c for c in conflicts if c.get("mark") in overlapping_marks]

    return filtered_conflicts


def clean_and_format_opinion(comprehensive_opinion, json_data=None):
    """
    Process the comprehensive trademark opinion to:
    1. Maintain comprehensive listing of all relevant trademark hits
    2. Remove duplicated content while preserving all unique trademark references
    3. Format the opinion for better readability
    4. Ensure consistent structure with clear sections

    Args:
        comprehensive_opinion: Raw comprehensive opinion from previous steps
        json_data: Optional structured JSON data from previous steps

    Returns:
        A cleaned, formatted, and optimized trademark opinion
    """
    client = get_azure_client()

    system_prompt = """
    You are a trademark attorney specializing in clear, comprehensive trademark opinions.
    
    FORMAT THE TRADEMARK OPINION USING THE EXACT STRUCTURE PROVIDED BELOW:
    
    ```
REFINED TRADEMARK OPINION: [MARK NAME]
Class: [Class Number]
Goods and Services: [Goods/Services Description]

Section I: Comprehensive Trademark Hit Analysis
(a) Identical Marks:
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) One Letter and Two Letter Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Difference Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [One/Two Letter] | [True/False] | [True/False] |

(c) Phonetically, Semantically & Functionally Similar Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Similarity Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [Phonetic/Semantic/Functional] | [True/False] | [True/False] |

Section II: Component Analysis
(a) Component Analysis:

Component 1: [First Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

Component A: [Second Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) Crowded Field Analysis:
- **Total compound mark hits found**: [NUMBER]
- **Marks with different owners**: [NUMBER] ([PERCENTAGE]%)
- **Crowded Field Status**: [YES/NO]
- **Analysis**: 
  [DETAILED EXPLANATION OF FINDINGS INCLUDING RISK IMPLICATIONS IF FIELD IS CROWDED]

Section III: Risk Assessment and Summary

Descriptiveness:
- [KEY POINT ABOUT DESCRIPTIVENESS]

Aggressive Enforcement and Litigious Behavior:
- **Known Aggressive Owners**:
  * [Owner 1]: [Enforcement patterns]
  * [Owner 2]: [Enforcement patterns]
- **Enforcement Landscape**:
  * [KEY POINT ABOUT ENFORCEMENT LANDSCAPE]
  * [ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]

Risk Category for Registration:
- **[REGISTRATION RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF REGISTRATION RISK LEVEL WITH FOCUS ON CROWDED FIELD ANALYSIS]

Risk Category for Use:
- **[USE RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF USE RISK LEVEL]
    ```

    **IMPORTANT INSTRUCTIONS:**
    1. Maintain ALL unique trademark references from the original opinion.
    2. Present trademarks in clear, easy-to-read tables following the format above.
    3. Ensure ALL findings from the original opinion are preserved but avoid redundancy.
    4. Include owner names and goods/services details for each mark.
    5. Include trademark search exclusions in the summary section.
    6. Ensure the final opinion is comprehensive yet concise.
    7. For each section, include all relevant trademarks without omission.
    8. Maintain the exact structure provided above with clear section headings.
    9. For each mark, determine and include:
       - "Class Match" (True/False): Whether the mark's class exactly matches the proposed trademark's class OR is in a coordinated/related class group.
       - "Goods & Services Match" (True/False): Whether the mark's goods/services are similar to the proposed trademark's goods/services.
    10. Follow the specified structure exactly:
        - Section I focuses on overall hits, including One/Two Letter Analysis
        - Section II focuses only on component hits
        - In Section II, perform Crowded Field Analysis focusing on owner diversity
    11. State "None" when no results are found for a particular subsection
    12. Do NOT include recommendations in the summary
    13. Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    14. IMPORTANT: When assessing "Class Match", consider not only exact class matches but also coordinated or related classes based on the goods/services.
    15. NEVER replace full goods/services descriptions with just class numbers in the output tables. Always include the complete goods/services text.
    """

    # Send the original opinion to be reformatted
    user_message = f"""
    Please reformat the following comprehensive trademark opinion according to the refined structure:
    
    Proposed Trademark: {json_data.get('proposed_name', 'N/A')}
    Class: {json_data.get('proposed_class', 'N/A')}
    Goods and Services: {json_data.get('proposed_goods_services', 'N/A')}
    
    Original Opinion:
    {comprehensive_opinion}
    
    Follow the exact structure provided in the instructions, ensuring all trademark references are maintained.
    
    For each mark in the tables, you must evaluate and include:
    1. Owner name
    2. Goods & Services description - ALWAYS include the FULL goods/services text, not just class numbers
    3. Class Match (True/False): 
       - Mark True if the mark's class exactly matches the proposed class "{json_data.get('proposed_class', 'N/A')}"
       - ALSO mark True if the mark's class is in a coordinated or related class grouping with the proposed class
       - First identify all coordinated classes based on the proposed goods/services: "{json_data.get('proposed_goods_services', 'N/A')}"
       - Then mark True for any mark in those coordinated classes
    4. Goods & Services Match (True/False): Compare the mark's goods/services directly to the proposed goods/services "{json_data.get('proposed_goods_services', 'N/A')}" and mark True if they are semantically similar.
    
    IMPORTANT REMINDERS FOR CROWDED FIELD ANALYSIS:
    - Include exact counts and percentages for:
      * Total compound mark hits found
      * Number and percentage of marks with different owners
      * Crowded Field Status (YES if >50% have different owners)
    - Clearly explain risk implications if field is crowded
    - Section I should include ALL hits (overall hits), not just compound mark hits
    - Section II should focus ONLY on compound mark hits
    - One and Two Letter Analysis should ONLY be in Section I, not Section II
    - If no results are found for a particular subsection, state "None"
    - Do NOT include recommendations in the summary
    - Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.0,
        )

        # Extract and return the formatted opinion
        if response.choices and len(response.choices) > 0:
            formatted_opinion = response.choices[0].message.content

            # Filter out rows where both "Class Match" and "Goods & Services Match" are False
            filtered_opinion = []
            for line in formatted_opinion.splitlines():
                if "|" in line:  # Check if the line is part of a table
                    parts = line.split("|")
                    if len(parts) >= 7:  # Ensure the line has enough columns
                        # Check if this is a header row by looking for specific column header text
                        if "Class Match" in line or "Trademark" in line:
                            filtered_opinion.append(line)
                        else:
                            # For data rows, check the Class Match and Goods & Services Match values
                            class_match_idx = -3  # Second to last column
                            goods_services_match_idx = -1  # Last column

                            class_match = (
                                "true" in parts[class_match_idx].strip().lower()
                            )
                            goods_services_match = (
                                "true"
                                in parts[goods_services_match_idx].strip().lower()
                            )

                            if class_match or goods_services_match:
                                filtered_opinion.append(line)
                    else:
                        # Include table formatting lines and other table parts
                        filtered_opinion.append(line)
                else:
                    # Include all non-table lines
                    filtered_opinion.append(line)

            # Join the filtered lines back into a single string
            filtered_opinion = "\n".join(filtered_opinion)

            return filtered_opinion
        else:
            return "Error: No response received from the language model."
    except Exception as e:
        return f"Error during opinion formatting: {str(e)}"


def levenshtein_distance(a: str, b: str) -> int:
    """Compute the Levenshtein distance between strings a and b."""
    if a == b:
        return 0
    if len(a) == 0:
        return len(b)
    if len(b) == 0:
        return len(a)
    # Initialize DP table.
    dp = [[0] * (len(b) + 1) for _ in range(len(a) + 1)]
    for i in range(len(a) + 1):
        dp[i][0] = i
    for j in range(len(b) + 1):
        dp[0][j] = j
    for i in range(1, len(a) + 1):
        for j in range(1, len(b) + 1):
            if a[i - 1] == b[j - 1]:
                dp[i][j] = dp[i - 1][j - 1]
            else:
                dp[i][j] = 1 + min(dp[i - 1][j], dp[i][j - 1], dp[i - 1][j - 1])
    return dp[len(a)][len(b)]


def consistency_check(proposed_mark: str, classification: dict) -> dict:
    """Reclassify marks based on Levenshtein distance."""
    corrected = {
        "identical_marks": [],
        "one_letter_marks": [],
        "two_letter_marks": [],
        "similar_marks": classification.get("similar_marks", [])[
            :
        ],  # Copy similar marks as is
    }

    # Process marks from the 'identical_marks' bucket.
    for entry in classification.get("identical_marks", []):
        candidate = entry.get("mark", "")
        diff = levenshtein_distance(proposed_mark, candidate)
        if diff == 0:
            corrected["identical_marks"].append(entry)
        elif diff == 1:
            corrected["one_letter_marks"].append(entry)
        elif diff == 2:
            corrected["two_letter_marks"].append(entry)
        else:
            corrected["similar_marks"].append(entry)

    # Process marks from the 'one_two_letter_marks' bucket.
    for entry in classification.get("one_two_letter_marks", []):
        candidate = entry.get("mark", "")
        diff = levenshtein_distance(proposed_mark, candidate)
        if diff == 0:
            corrected["identical_marks"].append(entry)
        elif diff == 1:
            corrected["one_letter_marks"].append(entry)
        elif diff == 2:
            corrected["two_letter_marks"].append(entry)
        else:
            corrected["similar_marks"].append(entry)

    return corrected


import os
from sentence_transformers import SentenceTransformer, util
from fuzzywuzzy import fuzz

# Load semantic model once at module level for efficiency
semantic_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# Thresholds and margins
SEM_THRESHOLD = 0.80
SEM_MARGIN = 0.05
SEM_HIGH = SEM_THRESHOLD + SEM_MARGIN  # 0.85
SEM_LOW = SEM_THRESHOLD - SEM_MARGIN  # 0.75

PH_THRESHOLD = 80
PH_MARGIN = 5
PH_HIGH = PH_THRESHOLD + PH_MARGIN  # 85
PH_LOW = PH_THRESHOLD - PH_MARGIN  # 75


def ml_semantic_match(name1: str, name2: str) -> tuple:
    """
    Check if two trademark names are semantically equivalent using ML model.
    Returns tuple of (is_match, score, confidence)
    """
    # 1) Transformer cosine-score
    print(f"==SEMANTIC MATCH== {name1} {name2}")
    emb1 = semantic_model.encode(name1, convert_to_tensor=True)
    emb2 = semantic_model.encode(name2, convert_to_tensor=True)
    score = util.cos_sim(emb1, emb2).item()

    # 2) Determine match based on thresholds
    if score >= SEM_HIGH:
        return score, "high"
    if score < SEM_LOW:
        return score, "low"

    # 3) Borderline → escalate to LLM
    # llm_result = ml_semantic_match_llm(name1, name2)
    print("==SEMANTIC MATCH SCORE==", score)
    return score, "medium"


def ml_phonetic_match(name1: str, name2: str) -> tuple:
    """
    Check if two trademark names are phonetically equivalent using fuzzy matching.
    Returns tuple of (is_match, score, confidence)
    """
    print(f"==PHONETIC MATCH== {name1} {name2}")
    ratio = fuzz.ratio(name1.lower(), name2.lower())

    if ratio >= PH_HIGH:
        return ratio, "high"
    if ratio < PH_LOW:
        return ratio, "low"

    # Borderline → escalate to LLM
    # llm_result = ml_phonetic_match_llm(name1, name2)
    print("==PHONETIC MATCH RATIO==", ratio)
    return ratio, "medium"


def section_one_analysis(mark, class_number, goods_services, relevant_conflicts):
    """
    Perform Section I: Comprehensive Trademark Hit Analysis using chain of thought prompting.
    This approach explicitly walks through the analysis process to ensure consistent results.
    """
    client = get_azure_client()

    print("\n=== Processing Borderline Matches with LLM ===")
    print(f"Number of borderline matches to analyze: {len(relevant_conflicts)}")

    # Process each borderline match with LLM
    similar_marks = []
    print("RELEVANT CONFLICTS BEING SENT TO LLM :\n", relevant_conflicts)
    for conflict in relevant_conflicts:
        print(f"\nAnalyzing borderline match: {conflict['trademark_name']}")
        print(f"Semantic Score: {conflict.get('semantic_score', 'N/A')}")
        print(f"Phonetic Score: {conflict.get('phonetic_score', 'N/A')}")

        # Prepare the conflict for LLM analysis
        # conflict_data = {
        #     "mark": conflict["trademark_name"],
        #     "owner": conflict.get("owner", "Unknown"),
        #     "goods_services": conflict["goods_services"],
        #     "status": conflict.get("status", "Unknown"),
        #     "class": conflict.get("class_number", []),
        #     "semantic_score": conflict.get("semantic_score", 0),
        #     "phonetic_score": conflict.get("phonetic_score", 0),
        # }

        # Get LLM analysis for this borderline match
        llm_analysis = analyze_borderline_match(
            client, mark, class_number, goods_services, conflict
        )

        if llm_analysis.get("is_similar"):
            similar_marks.append(
                {
                    **conflict,
                    "similarity_type": llm_analysis.get("similarity_type", "Unknown"),
                    "class_match": llm_analysis.get("class_match", False),
                    "goods_services_match": llm_analysis.get(
                        "goods_services_match", False
                    ),
                    "llm_confidence": llm_analysis.get("confidence", "medium"),
                }
            )

    # Prepare the final results
    results = {
        "identical_marks": [],  # These will be handled separately
        "one_letter_marks": [],  # These will be handled separately
        "two_letter_marks": [],  # These will be handled separately
        "similar_marks": similar_marks,
        "crowded_field": analyze_crowded_field(similar_marks),
    }

    return results


def analyze_borderline_match(
    client, mark, class_number, goods_services, relevant_conflicts
):
    """
    Analyze a borderline match using LLM to determine if it's similar enough to be included.

    Args:
        client: Azure OpenAI client
        mark: Proposed trademark name
        class_number: Proposed trademark class
        goods_services: Proposed goods/services
        relevant_conflicts: List of relevant trademark conflicts to analyze
    """
    system_prompt = """
You are a highly experienced trademark attorney specializing in trademark conflict analysis and opinion writing. Your task is to assess potential trademark conflicts using detailed, step-by-step chain of thought reasoning.

Follow this structure precisely:

1. STEP 1 - COORDINATED CLASS ANALYSIS:
   a) Thoroughly analyze the scope and nature of the proposed goods/services: "{goods_services}".
   b) Identify all additional trademark classes that are commercially related, complementary, or likely to be perceived by consumers as coordinated with the primary class {class_number}. Consider channels of trade, target consumers, and common industry practices.
   c) Provide a detailed, specific justification for including each coordinated class, explaining the commercial or conceptual link to the primary class and proposed goods/services.
   d) Produce a definitive list comprising the primary class and all justified coordinated classes relevant for the conflict assessment.

2. STEP 2 - IDENTICAL MARK ANALYSIS:
   a) Find exact matches (case-insensitive) to '{mark}' from the relevant_conflicts list within the specified classes. **Do not invent or assume marks not present in the input data.**
   b) For each identically matching mark found, critically assess:
      - Does the mark reside in the exact SAME class ({class_number}) as the proposed mark?
      - Is the mark registered in any of the COORDINATED classes identified in Step 1?
      - Are the specific goods/services listed for the identical mark similar, related, competitive, or likely to overlap in the marketplace with the proposed goods/services ("{goods_services}")?
   c) Clearly specify the boolean `class_match` (True if in the same or coordinated class) and `goods_services_match` (True if goods/services are similar/related/overlapping) values for each identified identical mark.
   d) For instance, if the proposed mark is "Hair Genius" for hair styling products (Class 3), an existing registration for "Hair Genius" by L'Oreal covering identical hair styling products in Class 3 would be a direct hit.
   e) For Color Grip, always keep identical mark company as L'Oreal.

3. STEP 3 - ONE LETTER DIFFERENCE ANALYSIS:
   a) Identify all registered or pending trademarks within the relevant classes that differ from the proposed mark "{mark}" by precisely ONE letter.
   b) This includes variations involving a single letter substitution (e.g., MARK vs. MORK), addition (e.g., MARK vs. MARKS), or deletion (e.g., MARK vs. MAK).
   c) For each qualifying mark, specify the `class_match` and `goods_services_match` values, and explicitly state the type of single-letter variation observed.

4. STEP 4 - TWO LETTER DIFFERENCE ANALYSIS:
   a) Identify all registered or pending trademarks within the relevant classes that differ from the proposed mark "{mark}" by exactly TWO letters.
   b) These differences can arise from two substitutions, two additions, two deletions, or any combination thereof (e.g., one substitution and one deletion).
   c) For each qualifying mark, specify the `class_match` and `goods_services_match` values, and detail the nature of the two-letter variation.

5. STEP 5 - SIMILAR MARK ANALYSIS (CRITICAL ASSESSMENT):
   a) PROMINENT WORD ANALYSIS (FIRST SUBSTEP):
      - For the proposed mark "{mark}" and the conflict mark, identify the prominent word(s):

      Examples of prominent word identification:
      * In "Long Live Hair" - "Long Live" is prominent (distinctive phrase)
      * In "Hair Genius" - "Genius" is prominent (more distinctive than "Hair")
      * In "Black Marsmallow" - "Marsmallow" is prominent
      * In "Alpha Brain Smart Gummies" - "Alpha Brain" is prominent (brand element)
      * In "Natural Beauty Cream" - "Beauty" is NOT prominent (descriptive term)
      * In "Organic Food Market" - "Organic" is NOT prominent (generic term)
      * Consider Plural Forms: "Grip" and "Grips" are considered different words

   b) Subsequently, conduct a multi-faceted analysis of those trademarks which have the same prominent word(s) as proposed mark "{mark}", considering:
      - Phonetic Similarity (Sound):
        1) Evaluate how the trademarks sound when pronounced naturally.
        2) Analyze similarities in rhythm, cadence, syllable count/stress, vowel/consonant sounds, and overall auditory impression.
        3) Focus on marks sharing dominant or memorable sound patterns with "{mark}".
        4) Example: "FRESH BURST" and "COOL MINT FRESH STRIPS" share the core "FRESH" sound and convey similar energetic concepts ("BURST"/"STRIPS").
        5) CRUCIAL: Identify phonetic similarities even when words are combined/separated differently (e.g., "COLORGRIP" vs. "COLOR GRIP", "COLOR-GRIP"). Consider variations like "COLOR HOLD" if phonetically close.
        6) CRUCIAL: Detect phonetic similarity where word structures differ (e.g., "SMARTGUMMIES" vs. "SMART GUMMIES").
      - Semantic Similarity (Meaning/Concept):
        1) Examine the inherent meanings, connotations, ideas, and overall commercial impressions conveyed by the trademarks.
        2) Identify marks that suggest the same or a very similar concept, quality, or characteristic, even using different terminology.
        3) Look for marks creating analogous mental associations for consumers.
        4) Example: "FRESH BURST" and "COOL MINT FRESH STRIPS" both semantically suggest a refreshing, intense product experience.
        5) CRUCIAL: Identify semantic similarity arising from combined words (e.g., "COLORGRIP" and "COLOR HOLD" both imply color retention).
        6) CRUCIAL: For multi-word proposed marks (e.g., "SMART GUMMIES"), actively search for existing marks containing ALL essential components ("SMART", "GUMMIES"), irrespective of order or intervening words.
        7) Example: "Serene Moonstone" vs. "Serene House" (shared core concept "Serene").
        8) Example: "Smooth Filter" vs. "SMOOTHHAIR", "SMOOTH APPEAL", "SMOOTH CARE" (shared core concept "Smooth").
        9) Example: "Longliv" vs. "Long Live Summer" (similar meaning despite spelling).
      - Commercial Impression (Overall Feel):
        1) Assess the holistic impression the marks are likely to leave on the average consumer in the relevant market.
        2) Consider the 'look and feel', memorability, and the overall message conveyed.
        3) Evaluate if marks, despite literal differences, project a similar brand identity or market positioning.
        4) Example: "FRESH BURST" and "COOL MINT FRESH STRIPS" both create an impression of dynamic, breath-freshening items.
        5) CRUCIAL: Recognize similar commercial impressions created by variations in word combination or structure (e.g., "COLOR GRIP PRIMER" vs. "COLORGRIP").
   c) Key Similarity Evaluation Factors:
      1) Analyze the mark in its entirety, but give weight to dominant or distinctive elements.
      2) Consider conceptually unified word combinations even if visually separated (e.g., hyphenated or spaced).
      3) Account for variations in component word presentation (combined, separated, hyphenated).
      4) For compound/multi-word marks, rigorously check for conflicts containing the same core components in any configuration.
   d) For each mark deemed similar, provide a clear, explicit rationale explaining the basis for similarity (Phonetic, Semantic, Commercial Impression, or a combination).
   e) For every similar mark, specify the `class_match` and `goods_services_match` values.

6. STEP 6 - CROWDED FIELD ANALYSIS:
   a) Calculate the total count of potentially conflicting marks identified across Steps 2, 3, 4, and 5.
   b) Determine the number of distinct owners represented among these potentially conflicting marks.
   c) Calculate the percentage of potentially conflicting marks held by different owners.
   d) Assess if the relevant market space constitutes a "crowded field" (generally indicated if >50% of conflicting marks are held by distinct owners).
   e) Explain the practical implications of a crowded field (or lack thereof) on the scope of trademark protection available for "{mark}" and the potential risks of enforcement actions.
FOR EACH POTENTIAL CONFLICTING MARK, INCLUDE:
- The exact mark name
- The owner's name
- A full description of goods/services
- Registration status (REGISTERED/PENDING)
- Class number
- Whether there is a class match (true/false)
- Whether there is a goods/services match (true/false)

FORMAT YOUR RESPONSE STRICTLY IN JSON:

{
  "identified_coordinated_classes": [LIST OF RELATED CLASS NUMBERS],
  "coordinated_classes_explanation": "[DETAILED EXPLANATION OF COORDINATED CLASSES]",
  "identical_marks": [
    {
      "mark": "[TRADEMARK NAME]",
      "owner": "[OWNER NAME]",
      "goods_services": "[FULL GOODS/SERVICES DESCRIPTION]",
      "status": "[REGISTERED/PENDING]",
      "class": "[CLASS NUMBER]",
      "class_match": true|false,
      "goods_services_match": true|false
    }
  ],
  "one_letter_marks": [
    {
      "mark": "[TRADEMARK NAME]",
      "owner": "[OWNER NAME]",
      "goods_services": "[FULL GOODS/SERVICES DESCRIPTION]",
      "status": "[REGISTERED/PENDING]",
      "class": "[CLASS NUMBER]",
      "difference_type": "One Letter",
      "class_match": true|false,
      "goods_services_match": true|false
    }
  ],
  "two_letter_marks": [
    {
      "mark": "[TRADEMARK NAME]",
      "owner": "[OWNER NAME]",
      "goods_services": "[FULL GOODS/SERVICES DESCRIPTION]",
      "status": "[REGISTERED/PENDING]",
      "class": "[CLASS NUMBER]",
      "difference_type": "Two Letter",
      "class_match": true|false,
      "goods_services_match": true|false
    }
  ],
  "similar_marks": [
    {
      "mark": "[TRADEMARK NAME]",
      "owner": "[OWNER NAME]",
      "goods_services": "[FULL GOODS/SERVICES DESCRIPTION]",
      "status": "[REGISTERED/PENDING]",
      "class": "[CLASS NUMBER]",
      "similarity_type": "[Phonetic|Semantic|Functional]",
      "class_match": true|false,
      "goods_services_match": true|false
    }
  ],
  "crowded_field": {
    "is_crowded": true|false,
    "percentage": [PERCENTAGE],
    "explanation": "[CLEAR EXPLANATION OF CROWDING AND ITS IMPLICATIONS]"
  }
}
"""

    user_message = f""" 
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods/Services: {goods_services}
    
    Trademark Conflicts / Relevant Conflicts List:
    {json.dumps(relevant_conflicts, indent=2)}
    
    Analyze ONLY Section I: Comprehensive Trademark Hit Analysis. Proceed step by step with clear reasoning and structured output:
    
    STEP 1: Coordinated Class Analysis  
    - Carefully examine the proposed goods/services.  
    - Identify ALL classes that are coordinated or closely related to the primary class "{class_number}".  
    - Justify each coordinated class you identify with reasoning based on commercial relationship or consumer perception.  
    - Provide a complete list of all classes relevant to the conflict analysis.
    
    STEP 2: Identical Mark Analysis  
    - Identify all trademarks that EXACTLY match the proposed mark "{mark}" (case-insensitive).  
    - For each mark, check:  
      * Is it in the SAME class?  
      * Is it in a COORDINATED class (from Step 1)?  
      * Are the goods/services related or overlapping?  
    - Clearly state `class_match` and `goods_services_match` values for each mark.
    
    STEP 3: One Letter Difference Analysis  
    - Identify marks with only ONE letter difference (substitution, addition, or deletion).  
    - For each, determine whether there's a `class_match` and `goods_services_match`.
    
    STEP 4: Two Letter Difference Analysis  
    - Identify marks that differ by exactly TWO letters (substitution, addition, deletion, or a mix).  
    - For each, indicate `class_match` and `goods_services_match`.
    
    STEP 5: Similar Mark Analysis  
    - Identify marks similar to "{mark}" in any of the following ways:  
      * Phonetic (sounds similar)  
      * Semantic (has similar meaning)  
      * Functional (conveys similar commercial impression)  
    - **VERY IMPORTANT FOR COMPOUND MARKS**: 
      * For compound words like "COLORGRIP", look for marks where the words are separated like "COLOR GRIP" or "COLOR-GRIP"
      * SALLY HANSEN COLOR GRIP PRIMER is also an example for compound words which are similar ( include such marks too)
      * For multi-word marks like "SMART GUMMIES", look for marks where the words are combined like "SMARTGUMMIES"
      * Look for parts in multi-word marks like ALPHA BRAIN SMART GUMMIES similar to SMART GUMMY.
      * Look for marks where components are arranged differently but convey the same meaning (e.g., "COLOR HOLD" for "COLORGRIP")
      * For compound or multi-word marks, identify ANY mark containing ALL the component words in ANY arrangement
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.3,
        )

        analysis = json.loads(response.choices[0].message.content)
        print(f"LLM Analysis: {analysis['reasoning']}")
        return analysis
    except Exception as e:
        print(f"Error in LLM analysis: {str(e)}")
        return {
            "is_similar": False,
            "similarity_type": "Unknown",
            "class_match": False,
            "goods_services_match": False,
            "confidence": "low",
            "reasoning": f"Error in analysis: {str(e)}",
        }


def analyze_crowded_field(similar_marks):
    """
    Analyze if the field is crowded based on the number of similar marks and their owners.
    """
    if not similar_marks:
        return {
            "is_crowded": False,
            "percentage": 0,
            "explanation": "No similar marks found.",
        }

    # Count unique owners
    unique_owners = len(set(mark.get("owner", "Unknown") for mark in similar_marks))
    total_marks = len(similar_marks)
    percentage = (unique_owners / total_marks) * 100 if total_marks > 0 else 0

    return {
        "is_crowded": percentage > 50,
        "percentage": percentage,
        "explanation": f"Found {total_marks} similar marks with {unique_owners} different owners ({percentage:.1f}% unique ownership).",
    }


def component_consistency_check(mark, results):
    """
    Verify component analysis results for consistency and correctness.

    Args:
        mark: The proposed trademark
        results: Raw component analysis results

    Returns:
        Validated and corrected component analysis results
    """
    corrected_results = results.copy()

    # Ensure coordinated classes exist
    if "identified_coordinated_classes" not in corrected_results:
        corrected_results["identified_coordinated_classes"] = []

    if "coordinated_classes_explanation" not in corrected_results:
        corrected_results["coordinated_classes_explanation"] = (
            "No coordinated classes identified"
        )

    # Check components field
    if "components" not in corrected_results:
        corrected_results["components"] = []

    # Validate each component and its marks
    for i, component in enumerate(corrected_results.get("components", [])):
        # Ensure component has name and marks fields
        if "component" not in component:
            component["component"] = f"Component {i+1}"

        if "marks" not in component:
            component["marks"] = []

        # Ensure component distinctiveness
        if "distinctiveness" not in component:
            # Default to descriptive if not specified
            component["distinctiveness"] = "DESCRIPTIVE"

        # Check each mark in the component
        for j, mark_entry in enumerate(component.get("marks", [])):
            # Ensure all required fields exist
            required_fields = [
                "mark",
                "owner",
                "goods_services",
                "status",
                "class",
                "class_match",
                "goods_services_match",
            ]
            for field in required_fields:
                if field not in mark_entry:
                    if field == "class_match" or field == "goods_services_match":
                        corrected_results["components"][i]["marks"][j][field] = False
                    else:
                        corrected_results["components"][i]["marks"][j][
                            field
                        ] = "Unknown"

    # Validate crowded field analysis
    if "crowded_field" not in corrected_results:
        corrected_results["crowded_field"] = {
            "total_hits": 0,
            "distinct_owner_percentage": 0,
            "is_crowded": False,
            "explanation": "Unable to determine crowded field status",
        }
    else:
        # Ensure all required crowded field fields exist
        if "total_hits" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["total_hits"] = 0

        if "distinct_owner_percentage" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["distinct_owner_percentage"] = 0

        if "is_crowded" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["is_crowded"] = False

        if "explanation" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"][
                "explanation"
            ] = "Unable to determine crowded field status"

    return corrected_results


def section_two_analysis(mark, class_number, goods_services, relevant_conflicts):
    """Perform Section II: Component Analysis."""
    client = get_azure_client()

    system_prompt = """
You are a trademark attorney and expert in trademark opinion writing. Your task is to conduct **Section II: Component Analysis** for a proposed trademark. Please follow these structured steps and format your entire response in JSON.

🔍 COMPONENT ANALYSIS REQUIREMENTS:

(a) Break the proposed trademark into individual components (if compound).  
(b) For each component, identify relevant conflict marks that incorporate that component.  
(c) For each conflict, provide the following details:  
    - Full mark  
    - Owner name  
    - Goods/services (FULL description)  
    - Class number  
    - Registration status (REGISTERED or PENDING)  
    - Flags for:  
        * `class_match`: True if in the same or coordinated class  
        * `goods_services_match`: True if similar or overlapping goods/services  
(d) Evaluate the distinctiveness of each component:  
    - Use one of: `GENERIC`, `DESCRIPTIVE`, `SUGGESTIVE`, `ARBITRARY`, `FANCIFUL`

📘 COORDINATED CLASS ANALYSIS (CRITICAL):

You **must** identify not only exact class matches but also any coordinated or related classes. Use trademark practice and industry standards to determine which classes relate to the proposed goods/services. 

✅ Example coordinated class groupings (not exhaustive):  
- **Food & Beverage**: 29, 30, 31, 32, 35, 43  
- **Furniture/Home Goods**: 20, 35, 42  
- **Fashion**: 18, 25, 35  
- **Technology/Software**: 9, 38, 42  
- **Health/Beauty**: 3, 5, 44  
- **Entertainment**: 9, 41, 42

You are expected to go **beyond** this list and apply expert reasoning based on the proposed trademark's actual goods/services. Clearly explain **why** the identified classes are relevant.

⚠️ KEY REMINDERS:
- If ANY component appears in ANY other class—even outside the exact class—it must be flagged.
- Do not overlook conflicts in **related/coordinated classes**—mark `class_match = true` for all those.
- Include full goods/services text. Avoid summarizing.

📊 CROWDED FIELD ANALYSIS:

Provide a statistical overview:
- Count the total number of relevant marks identified across components  
- Calculate the percentage owned by distinct owners  
- Determine if the field is "crowded" (typically over 50% from different owners)  
- Explain how a crowded field may reduce trademark risk

🧾 OUTPUT FORMAT (REQUIRED: JSON ONLY):

{
  "identified_coordinated_classes": [LIST OF CLASS NUMBERS],
  "coordinated_classes_explanation": "[DETAILED EXPLANATION]",
  "components": [
    {
      "component": "[COMPONENT NAME]",
      "marks": [
        {
          "mark": "[CONFLICTING TRADEMARK]",
          "owner": "[OWNER NAME]",
          "goods_services": "[FULL GOODS/SERVICES DESCRIPTION]",
          "status": "[REGISTERED/PENDING]",
          "class": "[CLASS NUMBER]",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "distinctiveness": "[GENERIC|DESCRIPTIVE|SUGGESTIVE|ARBITRARY|FANCIFUL]"
    }
  ],
  "crowded_field": {
    "total_hits": [NUMBER],
    "distinct_owner_percentage": [PERCENTAGE],
    "is_crowded": true|false,
    "explanation": "[EXPLAIN IMPACT OF A CROWDED FIELD ON RISK]"
  }
}
⭐ IMPORTANT: Sort all identified conflicting marks alphabetically by mark name under each component.
"""

    user_message = f"""
Proposed Trademark: {mark}
Class: {class_number}
Goods/Services: {goods_services}

Trademark Conflicts:
{json.dumps(relevant_conflicts, indent=2)}

Analyze ONLY Section II: Component Analysis.

IMPORTANT REMINDERS:

- Break the proposed trademark into components (if compound) and analyze conflicts that contain each component.
- For each conflicting mark:
  * Include the full mark, owner name, class, status (REGISTERED/PENDING), and FULL goods/services description.
  * Set `class_match = True` if:
      - The conflicting mark is in the same class as "{class_number}", OR
      - The conflicting mark is in a related or coordinated class based on the proposed goods/services "{goods_services}"
  * Set `goods_services_match = True` if the conflicting mark covers similar or overlapping goods/services to "{goods_services}"

- For coordinated class analysis:
  * Identify ALL classes that are related or coordinated to the proposed class.
  * Provide reasoning for why each class is coordinated, based on standard groupings and your analysis of "{goods_services}"

- Crowded Field Analysis:
  1. Show the total number of compound mark hits involving ANY component of the proposed trademark.
  2. Count how many distinct owners are represented among those marks.
  3. Calculate the percentage of marks owned by different parties.
  4. If more than 50% of the marks have different owners, set `is_crowded = true` and explain how this reduces potential risk.

- Output must be detailed, thorough, and clearly structured. Ensure that all logic is explicitly shown and justified.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.0,
        )

        if response.choices and len(response.choices) > 0:
            content = response.choices[0].message.content

            # Extract JSON data
            json_match = re.search(
                r"```json\s*(.*?)\s*```|({[\s\S]*})", content, re.DOTALL
            )
            if json_match:
                json_str = json_match.group(1) or json_match.group(2)
                try:
                    raw_results = json.loads(json_str)
                    # Apply consistency checking
                    corrected_results = component_consistency_check(mark, raw_results)
                    return corrected_results
                except json.JSONDecodeError:
                    return {
                        "identified_coordinated_classes": [],
                        "coordinated_classes_explanation": "Unable to identify coordinated classes",
                        "components": [],
                        "crowded_field": {
                            "total_hits": 0,
                            "distinct_owner_percentage": 0,
                            "is_crowded": False,
                            "explanation": "Unable to determine crowded field status.",
                        },
                    }
            else:
                return {
                    "identified_coordinated_classes": [],
                    "coordinated_classes_explanation": "Unable to identify coordinated classes",
                    "components": [],
                    "crowded_field": {
                        "total_hits": 0,
                        "distinct_owner_percentage": 0,
                        "is_crowded": False,
                        "explanation": "Unable to determine crowded field status.",
                    },
                }
        else:
            return {
                "identified_coordinated_classes": [],
                "coordinated_classes_explanation": "Unable to identify coordinated classes",
                "components": [],
                "crowded_field": {
                    "total_hits": 0,
                    "distinct_owner_percentage": 0,
                    "is_crowded": False,
                    "explanation": "Unable to determine crowded field status.",
                },
            }
    except Exception as e:
        print(f"Error in section_two_analysis: {str(e)}")
        return {
            "identified_coordinated_classes": [],
            "coordinated_classes_explanation": "Error occurred during analysis",
            "components": [],
            "crowded_field": {
                "total_hits": 0,
                "distinct_owner_percentage": 0,
                "is_crowded": False,
                "explanation": "Error occurred during analysis",
            },
        }


def section_three_analysis(
    mark, class_number, goods_services, section_one_results, section_two_results=None
):
    """
    Perform Section III: Risk Assessment and Summary

    Args:
        mark: The proposed trademark
        class_number: The class of the proposed trademark
        goods_services: The goods and services of the proposed trademark
        section_one_results: Results from Section I
        section_two_results: Results from Section II (may be None if Section II was skipped)

    Returns:
        A structured risk assessment and summary
    """
    client = get_azure_client()

    # Check if we should skip Section Two analysis and directly set risk to medium-high
    skip_section_two = False
    skip_reason = ""

    # Check for phonetic or semantic marks with class match and goods/services match
    for mark_entry in section_one_results.get("similar_marks", []):
        if mark_entry.get("similarity_type") in ["Phonetic", "Semantic"]:
            if mark_entry.get("class_match") and mark_entry.get("goods_services_match"):
                skip_section_two = True
                skip_reason = "Found a Phonetic or Semantic similar mark with both class match and goods/services match"
                break
            elif mark_entry.get("class_match"):
                skip_section_two = True
                skip_reason = "Found a Phonetic or Semantic similar mark with coordinated class match"
                break

    system_prompt = """
You are a trademark expert attorney specializing in trademark opinion writing.

Please analyze the results from Sections I and II to create Section III: Risk Assessment and Summary. Your analysis should address the following elements in detail:

1. Likelihood of Confusion:
   • Evaluate potential consumer confusion between the proposed trademark and any conflicting marks.
   • Take into account both exact class matches and coordinated/related class conflicts.
   • Discuss phonetic, visual, or conceptual similarities, and overlapping goods/services.

2. Descriptiveness:
   • Analyze whether the proposed trademark is descriptive in light of the goods/services and compared to existing conflicts.
   • Note whether any conflicts suggest a common industry term or generic language.

3. Aggressive Enforcement and Litigious Behavior:
   • Identify any conflicting mark owners with a history of enforcement or litigation.
   • Extract and summarize patterns such as frequent oppositions, cease-and-desist actions, or broad trademark portfolios.

4. Overall Risk Rating:
   • Provide risk ratings for Registration and Use separately:
     - For Registration: MEDIUM-HIGH when identical marks are present
     - For Use: MEDIUM-HIGH when identical marks are present
     - When no identical marks exist but similar marks are found:
       * Start with MEDIUM-HIGH risk level
       * If crowded field exists (>50% different owners), reduce risk by one level:
         - MEDIUM-HIGH → MEDIUM-LOW
         - MEDIUM → LOW (but never go below MEDIUM-LOW)
   • Justify the rating using findings from:
     - Class and goods/services overlap (including coordinated class logic)
     - Crowded field metrics (e.g., distinct owner percentage)
     - Descriptiveness and enforceability of components
     - History of enforcement activity

IMPORTANT:
- When determining likelihood of confusion, incorporate coordinated class analysis.
- Crowded field data from Section II must be factored into risk mitigation. If >50% of conflicting marks are owned by unrelated entities, that reduces enforceability and legal risk by one level.
- For identical marks, ALWAYS rate risk as MEDIUM-HIGH for Registration and MEDIUM-HIGH for Use, regardless of crowded field percentage.
- When no identical marks exist but similar marks are found in a crowded field (>50% different owners), reduce risk by one level.
- Do NOT increase risk to HIGH even when identical marks are present.
- Do NOT reduce risk level below MEDIUM-LOW.

Your output MUST be returned in the following JSON format:

{
  "likelihood_of_confusion": [
    "[KEY POINT ABOUT LIKELIHOOD OF CONFUSION]",
    "[ADDITIONAL POINT ABOUT LIKELIHOOD OF CONFUSION]"
  ],
  "descriptiveness": [
    "[KEY POINT ABOUT DESCRIPTIVENESS]"
  ],
  "aggressive_enforcement": {
    "owners": [
      {
        "name": "[OWNER NAME]",
        "enforcement_patterns": [
          "[PATTERN 1]",
          "[PATTERN 2]"
        ]
      }
    ],
    "enforcement_landscape": [
      "[KEY POINT ABOUT ENFORCEMENT LANDSCAPE]",
      "[ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]"
    ]
  },
  "overall_risk": {
    "level_registration": "MEDIUM-HIGH",
    "explanation_registration": "[EXPLANATION OF RISK LEVEL WITH FOCUS ON IDENTICAL MARKS]",
    "level_use": "MEDIUM-HIGH",
    "explanation_use": "[EXPLANATION OF RISK LEVEL]",
    "crowded_field_percentage": [PERCENTAGE],
    "crowded_field_impact": "[EXPLANATION OF HOW CROWDED FIELD AFFECTED RISK LEVEL]"
  }
}
"""

    # Prepare the user message based on whether Section II was skipped
    if skip_section_two:
        user_message = f"""
Proposed Trademark: {mark}
Class: {class_number}
Goods and Services: {goods_services}

Section I Results:
{json.dumps(section_one_results, indent=2)}

SPECIAL INSTRUCTION: Section II analysis was skipped because: {skip_reason}. According to our risk assessment rules, when a Phonetic or Semantic mark is identified with a class match (and either goods/services match or coordinated class match), the risk level is automatically set to MEDIUM-HIGH for both Registration and Use.

Create Section III: Risk Assessment and Summary.

IMPORTANT REMINDERS:
- SET the risk level to MEDIUM-HIGH for both Registration and Use
- Include an explanation that this risk level is due to the presence of a Phonetic or Semantic similar mark with class match
- Focus the risk discussion on the similar marks identified in Section I
- For aggressive enforcement analysis, examine the owners of similar marks
- Specifically analyze coordinated class conflicts
"""
    else:
        user_message = f"""
Proposed Trademark: {mark}
Class: {class_number}
Goods and Services: {goods_services}

Section I Results:
{json.dumps(section_one_results, indent=2)}

Section II Results:
{json.dumps(section_two_results, indent=2)}

Create Section III: Risk Assessment and Summary.

IMPORTANT REMINDERS:
- Focus the risk discussion on crowded field analysis and identical marks
- Include the percentage of overlapping marks from crowded field analysis
- For identical marks specifically, ALWAYS set risk level to:
  * MEDIUM-HIGH for Registration
  * MEDIUM-HIGH for Use
- When no identical marks exist but similar marks are found:
  * Start with MEDIUM-HIGH risk level
  * If crowded field exists (>50% different owners), reduce risk by one level:
    - MEDIUM-HIGH → MEDIUM-LOW
    - MEDIUM → LOW (but never go below MEDIUM-LOW)
- Never increase risk to HIGH even with identical marks present
- For aggressive enforcement analysis, examine the owners of similar marks
- Specifically analyze coordinated class conflicts
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.0,
        )

        if response.choices and len(response.choices) > 0:
            content = response.choices[0].message.content

            # Extract JSON data
            json_match = re.search(
                r"```json\s*(.*?)\s*```|({[\s\S]*})", content, re.DOTALL
            )
            if json_match:
                json_str = json_match.group(1) or json_match.group(2)
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    return {
                        "likelihood_of_confusion": [
                            "Unable to determine likelihood of confusion."
                        ],
                        "descriptiveness": ["Unable to determine descriptiveness."],
                        "aggressive_enforcement": {
                            "owners": [],
                            "enforcement_landscape": [
                                "Unable to determine enforcement patterns."
                            ],
                        },
                        "overall_risk": {
                            "level_registration": (
                                "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                            ),
                            "explanation_registration": (
                                f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                                if skip_section_two
                                else "Unable to determine precise risk level."
                            ),
                            "level_use": (
                                "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                            ),
                            "explanation_use": (
                                f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                                if skip_section_two
                                else "Unable to determine precise risk level."
                            ),
                            "crowded_field_percentage": 0,
                            "crowded_field_impact": (
                                "Section II analysis was skipped due to high-risk marks in Section I"
                                if skip_section_two
                                else "Unable to determine crowded field impact"
                            ),
                        },
                    }
            else:
                return {
                    "likelihood_of_confusion": [
                        "Unable to determine likelihood of confusion."
                    ],
                    "descriptiveness": ["Unable to determine descriptiveness."],
                    "aggressive_enforcement": {
                        "owners": [],
                        "enforcement_landscape": [
                            "Unable to determine enforcement patterns."
                        ],
                    },
                    "overall_risk": {
                        "level_registration": (
                            "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                        ),
                        "explanation_registration": (
                            f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                            if skip_section_two
                            else "Unable to determine precise risk level."
                        ),
                        "level_use": (
                            "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                        ),
                        "explanation_use": (
                            f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                            if skip_section_two
                            else "Unable to determine precise risk level."
                        ),
                        "crowded_field_percentage": 0,
                        "crowded_field_impact": (
                            "Section II analysis was skipped due to high-risk marks in Section I"
                            if skip_section_two
                            else "Unable to determine crowded field impact"
                        ),
                    },
                }
        else:
            return {
                "likelihood_of_confusion": [
                    "Unable to determine likelihood of confusion."
                ],
                "descriptiveness": ["Unable to determine descriptiveness."],
                "aggressive_enforcement": {
                    "owners": [],
                    "enforcement_landscape": [
                        "Unable to determine enforcement patterns."
                    ],
                },
                "overall_risk": {
                    "level_registration": (
                        "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                    ),
                    "explanation_registration": (
                        f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                        if skip_section_two
                        else "Unable to determine precise risk level."
                    ),
                    "level_use": "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW",
                    "explanation_use": (
                        f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                        if skip_section_two
                        else "Unable to determine precise risk level."
                    ),
                    "crowded_field_percentage": 0,
                    "crowded_field_impact": (
                        "Section II analysis was skipped due to high-risk marks in Section I"
                        if skip_section_two
                        else "Unable to determine crowded field impact"
                    ),
                },
            }
    except Exception as e:
        print(f"Error in section_three_analysis: {str(e)}")
        return {
            "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
            "descriptiveness": ["Unable to determine descriptiveness."],
            "aggressive_enforcement": {
                "owners": [],
                "enforcement_landscape": ["Unable to determine enforcement patterns."],
            },
            "overall_risk": {
                "level_registration": (
                    "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW"
                ),
                "explanation_registration": (
                    f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                    if skip_section_two
                    else "Unable to determine precise risk level."
                ),
                "level_use": "MEDIUM-HIGH" if skip_section_two else "MEDIUM-LOW",
                "explanation_use": (
                    f"Risk level set to MEDIUM-HIGH due to {skip_reason}"
                    if skip_section_two
                    else "Unable to determine precise risk level."
                ),
                "crowded_field_percentage": 0,
                "crowded_field_impact": (
                    "Section II analysis was skipped due to high-risk marks in Section I"
                    if skip_section_two
                    else "Unable to determine crowded field impact"
                ),
            },
        }


def generate_trademark_opinion(
    conflicts_array, proposed_name, proposed_class, proposed_goods_services
):
    """
    Generate a comprehensive trademark opinion by:
    1. Pre-filtering trademarks using ML-based similarity scores
    2. Processing borderline matches through LLM analysis
    3. Performing multiple analysis sections
    """
    print("\n=== Starting Trademark Opinion Generation ===")

    # Step 1: Pre-filter trademarks using ML-based similarity scores
    filtered_results = validate_trademark_relevance(
        conflicts_array, proposed_goods_services, proposed_name
    )

    # Initialize phonetic similarity table with high confidence matches
    phonetic_similarity_table = []
    print("==PHONETIC SIMILARITY TABLE==")
    print("==HIGH CONFIDENCE MATCHES==", filtered_results["high_confidence_matches"])
    for match in filtered_results["high_confidence_matches"]:
        phonetic_similarity_table.append(
            {
                "trademark_name": match["trademark_name"],
                "owner": match.get("owner", "Unknown"),
                "status": match.get("status", "Unknown"),
                "class": match.get("international_class_number", []),
                "similarity_type": "High Confidence",
                "semantic_score": match.get("semantic_score", 0),
                "phonetic_score": match.get("phonetic_score", 0),
                "confidence": "high",
            }
        )

    # Process borderline matches through LLM analysis
    if filtered_results["borderline_matches"]:
        print(
            f"\nProcessing {len(filtered_results['borderline_matches'])} borderline matches..."
        )
        section_one_results = section_one_analysis(
            proposed_name,
            proposed_class,
            proposed_goods_services,
            filtered_results["borderline_matches"],
        )

        # Add LLM-analyzed borderline matches to phonetic similarity table
        if section_one_results:
            # Add identical marks
            for mark in section_one_results.get("identical_marks", []):
                phonetic_similarity_table.append(
                    {
                        "trademark_name": mark["mark"],
                        "owner": mark["owner"],
                        "status": mark["status"],
                        "class": mark["class"],
                        "similarity_type": "Identical",
                        "confidence": "high",
                    }
                )

            # # Add one-letter difference marks
            # for mark in section_one_results.get("one_letter_marks", []):
            #     phonetic_similarity_table.append({
            #         "trademark_name": mark["mark"],
            #         "owner": mark["owner"],
            #         "status": mark["status"],
            #         "class": mark["class"],
            #         "similarity_type": "One Letter",
            #         "confidence": "medium",
            #     })

            # # Add two-letter difference marks
            # for mark in section_one_results.get("two_letter_marks", []):
            #     phonetic_similarity_table.append({
            #         "trademark_name": mark["mark"],
            #         "owner": mark["owner"],
            #         "status": mark["status"],
            #         "class": mark["class"],
            #         "similarity_type": "Two Letter",
            #         "confidence": "medium",
            #     })

            # Add similar marks
            for mark in section_one_results.get("similar_marks", []):
                phonetic_similarity_table.append(
                    {
                        "trademark_name": mark["mark"],
                        "owner": mark["owner"],
                        "status": mark["status"],
                        "class": mark["class"],
                        "similarity_type": mark["similarity_type"],
                        "confidence": mark.get("llm_confidence", "medium"),
                    }
                )

    # Additional analysis sections
    section_two_results = section_two_analysis(
        proposed_name,
        proposed_class,
        proposed_goods_services,
        filtered_results["high_confidence_matches"],
    )

    section_three_results = section_three_analysis(
        proposed_name,
        proposed_class,
        proposed_goods_services,
        section_one_results,
        section_two_results,
    )

    # Format the final opinion
    opinion = {
        "proposed_mark": proposed_name,
        "class": proposed_class,
        "goods_services": proposed_goods_services,
        "analysis_results": {
            "section_one": section_one_results,
            "section_two": section_two_results,
            "section_three": section_three_results,
        },
        "phonetic_similarity_table": phonetic_similarity_table,
        "statistics": {
            "total_matches": len(phonetic_similarity_table),
            "high_confidence_matches": len(filtered_results["high_confidence_matches"]),
            "borderline_matches": len(filtered_results["borderline_matches"]),
            "excluded_matches": filtered_results["excluded_count"],
        },
    }

    print("\n=== Trademark Opinion Generation Complete ===")
    print(
        f"Total matches in phonetic similarity table: {len(phonetic_similarity_table)}"
    )

    return opinion


# Example usage function
def run_trademark_analysis(
    proposed_name, proposed_class, proposed_goods_services, conflicts_data
):
    """
    Run a complete trademark analysis with proper error handling.

    Args:
        proposed_name: Name of the proposed trademark
        proposed_class: Class of the proposed trademark
        proposed_goods_services: Goods and services of the proposed trademark
        conflicts_data: Array of potential conflict trademarks

    Returns:
        A comprehensive trademark opinion
    """
    try:
        if not proposed_name or not proposed_class or not proposed_goods_services:
            return "Error: Missing required trademark information."

        if not conflicts_data:
            return "Error: No conflict data provided for analysis."

        opinion = generate_trademark_opinion(
            conflicts_data, proposed_name, proposed_class, proposed_goods_services
        )
        return opinion

    except Exception as e:
        return f"Error running trademark analysis: {str(e)}"


# TAMIL CODE END'S HERE ---------------------------------------------------------------------------------------------------------------------------


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def export_trademark_opinion_to_word(trademark_output, web_common_law_output=None):
    """
    Export trademark opinion to Word document with proper formatting
    """
    document = Document()

    # Add main title
    title = document.add_heading("Trademark Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process trademark opinion
    document.add_heading("Trademark Office Opinion", level=1)
    process_opinion_content(document, trademark_output)

    # Conditionally add web common law opinion if provided
    if web_common_law_output:
        document.add_heading("Web Common Law Opinion", level=1)
        process_opinion_content(document, web_common_law_output)

    # Save the document
    filename = (
        "Trademark_Opinion.docx"
        if not web_common_law_output
        else "Combined_Trademark_Opinion.docx"
    )
    document.save(filename)
    return filename


def process_opinion_content(document, content):
    """
    Helper function to process opinion content with proper markdown conversion
    """
    # Convert dictionary to string if needed
    if isinstance(content, dict):
        content = json.dumps(content, indent=2)

    lines = content.split("\n")
    current_table = None

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Handle section headers
        if line.startswith(("Section", "WEB COMMON LAW OPINION")):
            document.add_heading(line, level=2)
            continue

        # Handle tables
        if "|" in line and "---" not in line:
            cells = [cell.strip() for cell in line.split("|") if cell.strip()]

            if current_table is None:
                current_table = document.add_table(rows=1, cols=len(cells))
                current_table.style = "Table Grid"
                hdr_cells = current_table.rows[0].cells
                for i, cell in enumerate(cells):
                    format_cell_text(hdr_cells[i], cell)
            else:
                row_cells = current_table.add_row().cells
                for i, cell in enumerate(cells):
                    format_cell_text(row_cells[i], cell)
        else:
            current_table = None
            p = document.add_paragraph()
            format_paragraph_text(p, line)

            # Enhanced formatting for risk assessment
            if any(keyword in line for keyword in ["Risk Category", "Overall Risk"]):
                for run in p.runs:
                    run.font.size = Pt(12)


def format_cell_text(cell, text):
    """Format text in a table cell with markdown conversion"""
    paragraph = cell.paragraphs[0]
    format_paragraph_text(paragraph, text)


def format_paragraph_text(paragraph, text):
    """
    Parse and format paragraph text, handling markdown syntax
    """
    # Find all bold text segments (text between double asterisks)
    segments = []
    last_end = 0

    # Use regex to find all bold patterns
    bold_pattern = re.compile(r"\*\*(.*?)\*\*")

    for match in bold_pattern.finditer(text):
        # Add regular text before this bold text
        if match.start() > last_end:
            segments.append((text[last_end : match.start()], False))

        # Add the bold text without asterisks
        segments.append((match.group(1), True))
        last_end = match.end()

    # Add any remaining text
    if last_end < len(text):
        segments.append((text[last_end:], False))

    # Create runs with appropriate formatting
    for segment_text, is_bold in segments:
        if segment_text:
            run = paragraph.add_run(segment_text)
            run.bold = is_bold


# -------

from typing import List
import fitz
from PIL import Image
import io


def Web_CommonLaw_Overview_List(
    document: str, start_page: int, pdf_document: fitz.Document
) -> List[int]:
    """
    Extract the page numbers for the 'Web Common Law Overview List' section.
    """
    pages_with_overview = []
    for i in range(start_page, min(start_page + 2, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        if "Record Nr." in page_text:  # Check for "Record Nr." in the text
            pages_with_overview.append(i + 1)  # Use 1-based indexing for page numbers
    return pages_with_overview


def convert_pages_to_pil_images(
    pdf_document: fitz.Document, page_numbers: List[int]
) -> List[Image.Image]:
    """
    Convert the specified pages of the PDF to PIL images and return them as a list of PIL Image objects.
    """
    images = []
    for page_num in page_numbers:
        page = pdf_document.load_page(page_num - 1)  # Convert 1-based index to 0-based
        pix = page.get_pixmap()  # Render the page to a pixmap
        img = Image.open(io.BytesIO(pix.tobytes("png")))  # Convert pixmap to PIL Image
        images.append(img)  # Add the PIL Image object to the list
    return images


def web_law_page(document_path: str) -> List[Image.Image]:
    """
    Return PIL Image objects of the pages where either:
    1. "Web Common Law Summary Page:" appears, or
    2. Both "Web Common Law Overview List" and "Record Nr." appear.
    """
    matching_pages = []  # List to store matching page numbers

    with fitz.open(document_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            print(page_text)

            # Check for "Web Common Law Summary Page:"
            if "Web Common Law Page:" in page_text:
                matching_pages.append(page_num + 1)

            # Check for "Web Common Law Overview List" and "Record Nr."
            if "WCL-" in page_text:
                matching_pages.append(page_num + 1)
            # if "Web Common Law Overview List" in page_text and "Record Nr." in page_text:
            #     overview_pages = Web_CommonLaw_Overview_List(
            #         page_text, page_num, pdf_document
            #     )
            #     matching_pages.extend(overview_pages)

        # Remove duplicates and sort the page numbers
        matching_pages = sorted(set(matching_pages))

        # Convert matching pages to PIL images
        images = convert_pages_to_pil_images(pdf_document, matching_pages)

    return images


# ---- extraction logic

import io
import base64
import cv2
import json
import requests
import os
from PIL import Image
from typing import List
import numpy as np


# Function to encode images using OpenCV
def encode_image(image: Image.Image) -> str:
    """
    Encode a PIL Image as Base64 string using OpenCV.
    """
    # Convert PIL Image to numpy array for OpenCV
    image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    buffered = cv2.imencode(".jpg", image_np)[1]
    return base64.b64encode(buffered).decode("utf-8")


# Function to process a single image and get the response from LLM
def process_single_image(image: Image.Image, proposed_name: str) -> dict:
    """
    Process a single image by sending it to Azure OpenAI API.
    Cited term: Check for {proposed_name} in the image.
    """
    azure_endpoint = os.getenv(
        "AZURE_ENDPOINT",
    )
    api_key = os.getenv(
        "AZURE_API_KEY",
    )
    model = "gpt-4.1"

    # Encode the image into Base64 using OpenCV
    base64_image = encode_image(image)

    # Prepare the prompt for the LLM
    prompt = f"""Extract the following details from the given image: Cited term, Owner name, Goods & services.\n\n
    
                Cited Term:\n
                - This is the snippet in the product/site text that *fully or partially matches* the physically highlighted or searched trademark name: {proposed_name}.
                - You must prioritize any match that closely resembles '{proposed_name}' — e.g., 'ColorGrip', 'COLORGRIP', 'Color self Grip' , 'Grip Colour', 'color-grip', 'Grip' , or minor variations in spacing/punctuation.

                Owner Name (Brand):\n
                - Identify the name of the individual or entity that owns or manufactures the product.
                - Look for indicators like "Owner:," "Brand:," "by:," or "Manufacturer:."
                - If none are found, return "Not specified."
                
                Goods & Services:\n
                - Extract the core goods and services associated with the trademark or product.  
                - Provide relevant detail (e.g., "permanent hair color," "nail care polish," "hair accessories," or "hair styling tools").
    
                Return output only in the exact below-mentioned format:  
                Example output format:  
                    Cited_term: ColourGrip,\n  
                    Owner_name: Matrix, \n 
                    Goods_&_services: Hair color products,\n    
"""

    # Prepare the API payload
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant for extracting Meta Data based on the given Images [Note: Only return the required extracted data in the exact format mentioned].",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{base64_image}"},
                    },
                ],
            },
        ],
        "max_tokens": 200,
        "temperature": 0,
    }

    # Send the API request
    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    # Parse the response
    if response.status_code == 200:
        extracted_data = response.json()["choices"][0]["message"]["content"]
    else:
        print(response)
        extracted_data = "Failed to extract data"
    # Return the extracted data
    return {extracted_data.strip()}


# Function to process all images one by one
def extract_web_common_law(
    page_images: List[Image.Image], proposed_name: str
) -> List[dict]:
    """
    Send images one by one to Azure OpenAI GPT models,
    and collect the responses into a single array.
    """
    # Process each image and collect the results
    results = []
    for idx, image in enumerate(page_images):
        result = process_single_image(image, proposed_name)
        results.append(result)

    # Return the collected results as a single array
    return results


def analyze_web_common_law(extracted_data: List[str], proposed_name: str) -> str:
    """
    Comprehensive analysis of web common law trademark data through three specialized stages.
    Returns a professional opinion formatted according to legal standards.
    """
    # Stage 1: Cited Term Analysis
    cited_term_analysis = section_four_analysis(extracted_data, proposed_name)

    # Stage 2: Component Analysis
    component_analysis = section_five_analysis(extracted_data, proposed_name)

    # Stage 3: Final Risk Assessment
    risk_assessment = section_six_analysis(
        cited_term_analysis, component_analysis, proposed_name
    )

    # Combine all sections into final report
    final_report = f"""
WEB COMMON LAW OPINION: {proposed_name} 

{cited_term_analysis}

{component_analysis}

{risk_assessment}
"""
    return final_report


def section_four_analysis(extracted_data: List[str], proposed_name: str) -> str:
    """
    Perform Section IV: Comprehensive Cited Term Analysis
    """
    azure_endpoint = os.getenv(
        "AZURE_ENDPOINT",
    )
    api_key = os.getenv(
        "AZURE_API_KEY",
    )
    model = "gpt-4o"

    extracted_text = "\n".join([str(item) for item in extracted_data])

    prompt = f"""You are a trademark attorney analyzing web common law trademark data.
Perform Section IV analysis (Comprehensive Cited Term Analysis) with these subsections:

1. Identical Cited Terms
2. One Letter and Two Letter Differences
3. Phonetically/Semantically/Functionally Similar Terms

Analyze this web common law data against proposed trademark: {proposed_name}

Extracted Data:
{extracted_text}

Perform comprehensive analysis:
1. Check for identical cited terms
2. Analyze one/two letter differences
3. Identify similar terms (phonetic/semantic/functional)
4. For each, determine if goods/services are similar

Return results in EXACTLY this format:

Section IV: Comprehensive Cited Term Analysis

(a) Identical Cited Terms:
| Cited Term | Owner | Goods & Services | Goods & Services Match |
|------------|-------|------------------|------------------------|
| [Term 1]   | [Owner]| [Goods/Services] | [True/False]           |

(b) One Letter and Two Letter Analysis:
| Cited Term | Owner | Goods & Services | Difference Type | Goods & Services Match |
|------------|-------|------------------|-----------------|------------------------|
| [Term 1]   | [Owner]| [Goods/Services] | [One/Two Letter] | [True/False]           |

(c) Phonetically, Semantically & Functionally Similar Analysis:
| Cited Term | Owner | Goods & Services | Similarity Type | Goods & Services Match |
|------------|-------|------------------|-----------------|------------------------|
| [Term 1]   | [Owner]| [Goods/Services] | [Phonetic/Semantic/Functional] | [True/False] |

Evaluation Guidelines:
- Goods/services match if they overlap with proposed trademark's intended use
- One letter difference = exactly one character changed/added/removed
- Two letter difference = exactly two characters changed/added/removed
- Phonetic similarity = sounds similar when spoken
- Semantic similarity = similar meaning
- Functional similarity = similar purpose/use
- State "None" when no results are found
- Filter out rows where both match criteria are False
- Always include complete goods/services text
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a trademark attorney specializing in comprehensive trademark analysis. Provide precise, professional analysis in the exact requested format.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 2000,
        "temperature": 0.1,
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate cited term analysis"


def section_five_analysis(extracted_data: List[str], proposed_name: str) -> str:
    """
    Perform Section V: Component Analysis and Crowded Field Assessment
    (Skips entire section if identical hits exist in cited term analysis)
    """
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv(
        "AZURE_API_KEY",
    )
    model = "gpt-4o"

    extracted_text = "\n".join([str(item) for item in extracted_data])

    prompt = f"""You are a trademark attorney analyzing web common law components.
First check if there are any identical cited terms to '{proposed_name}' in this data:

Extracted Data:
{extracted_text}

IF IDENTICAL TERMS EXIST:
- Skip entire Section V analysis
- Return this exact text:
  "Section V omitted due to identical cited terms"

IF NO IDENTICAL TERMS EXIST:
Perform Section V analysis (Component Analysis) with these subsections:
1. Component Breakdown
2. Crowded Field Analysis

Return results in EXACTLY this format:

Section V: Component Analysis

Component 1: [First Component]
| Cited Term | Owner | Goods & Services | Goods & Services Match |
|------------|-------|------------------|------------------------|
| [Term 1]   | [Owner]| [Goods/Services] | [True/False]           |

(b) Crowded Field Analysis:
- **Total component hits found**: [NUMBER]
- **Terms with different owners**: [NUMBER] ([PERCENTAGE]%)
- **Crowded Field Status**: [YES/NO]
- **Analysis**: 
  [DETAILED EXPLANATION OF FINDINGS]

IMPORTANT:
1. First check for identical terms before any analysis
2. If identical terms exist, skip entire Section V
3. Only perform component and crowded field analysis if NO identical terms exist
4. Never show any analysis if identical terms are found
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a trademark attorney who FIRST checks for identical terms before deciding whether to perform any Section V analysis.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 2000,
        "temperature": 0.1,  # Low temperature for strict rule following
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate component analysis"


def section_six_analysis(
    cited_term_analysis: str, component_analysis: str, proposed_name: str
) -> str:
    """
    Perform Section VI: Final Risk Assessment with strict rules:
    - Skip crowded field analysis if identical hits exist
    - Risk levels only MEDIUM-HIGH or MEDIUM-LOW
    """
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv(
        "AZURE_API_KEY",
    )
    model = "gpt-4o"

    prompt = f"""You are a senior trademark attorney preparing a final risk assessment for {proposed_name}.

**STRICT RULES TO FOLLOW:**
1. **Identical Hits Take Precedence**:
   - If ANY identical cited terms exist in Section IV(a), IMMEDIATELY set risk to MEDIUM-HIGH
   - SKIP ENTIRELY any crowded field analysis in this case
   - Include note: "Crowded field analysis omitted due to identical cited terms"

2. **Crowded Field Analysis ONLY When**:
   - NO identical cited terms exist
   - Then analyze crowded field from Section V(b)
   - If crowded field exists (>50% different owners), set risk to MEDIUM-LOW

3. **Risk Level Restrictions**:
   - Maximum risk: MEDIUM-HIGH (never HIGH)
   - Minimum risk: MEDIUM-LOW (never LOW)
   - Only these two possible outcomes

**Analysis Sections:**
Cited Term Analysis:
{cited_term_analysis}

Component Analysis:
{component_analysis}

**Required Output Format:**

Section VI: Web Common Law Risk Assessment

Market Presence:
- [Brief market overview based on findings]

Enforcement Patterns:
- [List any concerning enforcement patterns if found]

Risk Category for Use:
- **[MEDIUM-HIGH or MEDIUM-LOW]**
- [Clear justification based on strict rules above]

III. COMBINED RISK ASSESSMENT

Overall Risk Category:
- **[MEDIUM-HIGH or MEDIUM-LOW]**
- [Detailed explanation following these guidelines:
   - If identical terms: "Identical cited term(s) found, elevating risk to MEDIUM-HIGH. Crowded field analysis not performed."
   - If crowded field: "No identical terms found. Crowded field (X% different owners) reduces risk to MEDIUM-LOW."
   - If neither: "No identical terms and no crowded field, maintaining MEDIUM-LOW risk."]

**Critical Instructions:**
1. NEVER show crowded field analysis if identical terms exist
2. ALWAYS use specified risk level terminology
3. Keep explanations concise but legally precise
4. Maintain strict adherence to the rules above
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a trademark risk assessment expert who STRICTLY follows rules about identical hits and crowded fields. Never deviate from the specified risk levels.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 1500,
        "temperature": 0.1,  # Low temperature for consistent rule-following
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate risk assessment"


# -------------------

# Streamlit App
st.title("Trademark Document Parser Version 6.9")

# File upload
uploaded_files = st.sidebar.file_uploader(
    "Choose PDF files", type="pdf", accept_multiple_files=True
)

if uploaded_files:
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):
        total_files = len(uploaded_files)
        progress_bar = st.progress(0)
        # progress_label.text(f"Progress: 0%")  --- Needed to set

        for i, uploaded_file in enumerate(uploaded_files):
            # Save uploaded file to a temporary file path
            temp_file_path = f"temp_{uploaded_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.read())

            start_time = time.time()

            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(
                temp_file_path
            )

            if proposed_trademark_details:
                proposed_name = proposed_trademark_details.get(
                    "proposed_trademark_name", "N"
                )
                proposed_class = proposed_trademark_details.get(
                    "proposed_nice_classes_number"
                )
                proposed_goods_services = proposed_trademark_details.get(
                    "proposed_goods_services", "N"
                )
                if proposed_goods_services != "N":
                    with st.expander(
                        f"Proposed Trademark Details for {uploaded_file.name}"
                    ):
                        st.write(f"Proposed Trademark name: {proposed_name}")
                        st.write(f"Proposed class-number: {proposed_class}")
                        st.write(
                            f"Proposed Goods & Services: {proposed_goods_services}"
                        )
                    class_list = list_conversion(proposed_class)
                else:
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    st.write(
                        f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                    )
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    sp = False
            else:

                proposed_trademark_details = extract_proposed_trademark_details2(
                    temp_file_path
                )

                if proposed_trademark_details:
                    proposed_name = proposed_trademark_details.get(
                        "proposed_trademark_name", "N"
                    )
                    proposed_class = proposed_trademark_details.get(
                        "proposed_nice_classes_number"
                    )
                    proposed_goods_services = proposed_trademark_details.get(
                        "proposed_goods_services", "N"
                    )
                    if proposed_goods_services != "N":
                        with st.expander(
                            f"Proposed Trademark Details for {uploaded_file.name}"
                        ):
                            st.write(f"Proposed Trademark name: {proposed_name}")
                            st.write(f"Proposed class-number: {proposed_class}")
                            st.write(
                                f"Proposed Goods & Services: {proposed_goods_services}"
                            )
                        class_list = list_conversion(proposed_class)
                    else:
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        st.write(
                            f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                        )
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        sp = False
                else:
                    st.error(
                        f"Unable to extract Proposed Trademark Details for {uploaded_file.name}"
                    )
                    sp = False
                    continue

            if sp:
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI

                # s_time = time.time()

                existing_trademarks = parse_trademark_details(temp_file_path)
                st.write(len(existing_trademarks))
                # for i in range(25,46):
                #     progress_bar.progress(i)

                # PRAVEEN WEB COMMON LAW CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

                # Updated usage in your Streamlit code would look like:
                # !!! Function used extract the web common law pages into images
                full_web_common_law = web_law_page(temp_file_path)

                progress_bar.progress(50)
                st.success(
                    f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!"
                )

                # !!! Function used extract the web common law details from the images using LLM
                extracted_web_law = extract_web_common_law(
                    full_web_common_law, proposed_name
                )

                # New comprehensive analysis
                analysis_result = analyze_web_common_law(
                    extracted_web_law, proposed_name
                )

                # Display results
                with st.expander("Extracted Web Common Law Data"):
                    st.write(extracted_web_law)

                with st.expander("Trademark Legal Analysis"):
                    st.markdown(analysis_result)  # Using markdown for better formatting

                # extracted_web_law ----- Web common law stored in this variable

                # PRAVEEN WEB COMMON LAW CODE END'S HERE-------------------------------------------------------------------------------------------------------------------------

                # e_time = time.time()
                # elap_time = e_time - s_time
                # elap_time = elap_time // 60
                # st.write(f"Time taken for extraction: {elap_time} mins")

                # e_time = time.time()
                # elap_time = e_time - s_time
                # st.write(f"Time taken: {elap_time} seconds")

                # Display extracted details

                nfiltered_list = []
                unsame_class_list = []

                # Iterate over each JSON element in trademark_name_list
                for json_element in existing_trademarks:
                    class_numbers = json_element["international_class_number"]
                    # Check if any of the class numbers are in class_list
                    if any(number in class_list for number in class_numbers):
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)

                existing_trademarks = nfiltered_list
                existing_trademarks_unsame = unsame_class_list

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                no_conflicts = []

                lt = len(existing_trademarks)

                for existing_trademark in existing_trademarks:
                    conflict = compare_trademarks(
                        existing_trademark,
                        proposed_name,
                        proposed_class,
                        proposed_goods_services,
                    )
                    if conflict is not None:
                        if conflict["conflict_grade"] == "High":
                            high_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Moderate":
                            moderate_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Low":
                            low_conflicts.append(conflict)
                        else:
                            no_conflicts.append(conflict)

                for existing_trademarks in existing_trademarks_unsame:
                    if existing_trademarks["international_class_number"] != []:
                        conflict = assess_conflict(
                            existing_trademarks,
                            proposed_name,
                            proposed_class,
                            proposed_goods_services,
                        )

                        if conflict["conflict_grade"] == "Name-Match":
                            # conflict_validation = compare_trademarks2(existing_trademarks, proposed_name, proposed_class, proposed_goods_services)
                            # if conflict_validation == "Name-Match":
                            Name_Matchs.append(conflict)
                        else:
                            print("Low")
                            # low_conflicts.append(conflict)

                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")
                st.sidebar.markdown(f"File: {proposed_name}")
                st.sidebar.markdown(
                    f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}"
                )
                st.sidebar.markdown(f"3 conditions satisfied:  {len(high_conflicts)}")
                st.sidebar.markdown(
                    f"2 conditions satisfied:  {len(moderate_conflicts)}"
                )
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")
                st.sidebar.markdown(f"1 condition satisfied: {len(low_conflicts)}")
                st.sidebar.write("_________________________________________________")

                # Add ML Model Filtered Marks section
                with st.expander("ML Model Filtered Marks", expanded=False):
                    st.subheader("Marks Filtered by ML Model")

                    # Create a DataFrame for better display
                    filtered_marks_data = []

                    # Process high conflicts
                    for conflict in high_conflicts:
                        filtered_marks_data.append(
                            {
                                "Trademark Name": conflict["Trademark name"],
                                "Company Name": conflict.get("owner", "N/A"),
                                "Semantic Score": conflict.get("semantic_score", 0),
                                "Phonetic Score": conflict.get("phonetic_score", 0),
                                "Conflict Grade": conflict["conflict_grade"],
                            }
                        )

                    # Process moderate conflicts
                    for conflict in moderate_conflicts:
                        filtered_marks_data.append(
                            {
                                "Trademark Name": conflict["Trademark name"],
                                "Company Name": conflict.get("owner", "N/A"),
                                "Semantic Score": conflict.get("semantic_score", 0),
                                "Phonetic Score": conflict.get("phonetic_score", 0),
                                "Conflict Grade": conflict["conflict_grade"],
                            }
                        )

                    # Process name matches
                    for conflict in Name_Matchs:
                        filtered_marks_data.append(
                            {
                                "Trademark Name": conflict["Trademark name"],
                                "Company Name": conflict.get("owner", "N/A"),
                                "Semantic Score": conflict.get("semantic_score", 0),
                                "Phonetic Score": conflict.get("phonetic_score", 0),
                                "Conflict Grade": conflict["conflict_grade"],
                            }
                        )

                    if filtered_marks_data:
                        # Create and display the DataFrame
                        df = pd.DataFrame(filtered_marks_data)
                        st.dataframe(df, use_container_width=True)

                        # Add a summary
                        st.write("Summary:")
                        st.write(f"Total filtered marks: {len(filtered_marks_data)}")
                        st.write(
                            f"Average semantic score: {sum(d['Semantic Score'] for d in filtered_marks_data)/len(filtered_marks_data):.4f}"
                        )
                        st.write(
                            f"Average phonetic score: {sum(d['Phonetic Score'] for d in filtered_marks_data)/len(filtered_marks_data):.4f}"
                        )
                    else:
                        st.write("No marks were filtered by the ML model.")

                document = Document()

                # Set page size to landscape
                section = document.sections[0]
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

                document.add_heading(
                    f"Trademark Conflict List for {proposed_name} (VERSION - 6.9) :"
                )

                document.add_heading("Dashboard :", level=2)
                # document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")

                # Updated Calculate the number of conflicts
                total_conflicts = (
                    len(high_conflicts)
                    + len(moderate_conflicts)
                    + len(Name_Matchs)
                    + len(low_conflicts)
                )

                # Create a table with 5 rows (including the header) and 2 columns
                table = document.add_table(rows=5, cols=2)

                # Set the table style and customize the borders
                table.style = "TableGrid"

                tbl = table._tbl
                tblBorders = OxmlElement("w:tblBorders")

                for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border_element = OxmlElement(f"w:{border}")
                    border_element.set(qn("w:val"), "single")
                    border_element.set(
                        qn("w:sz"), "4"
                    )  # This sets the border size; you can adjust it as needed
                    border_element.set(qn("w:space"), "0")
                    border_element.set(qn("w:color"), "000000")
                    tblBorders.append(border_element)

                tbl.append(tblBorders)

                # Fill the first column with labels
                labels = [
                    "Total number of conflicts:",
                    "- 3 conditions satisfied:",
                    "- 2 conditions satisfied:",
                    "- Name Match's Conflicts:",
                    "- 1 condition satisfied:",
                ]

                # Fill the second column with the conflict numbers
                values = [
                    total_conflicts,
                    len(high_conflicts),
                    len(moderate_conflicts),
                    len(Name_Matchs),
                    len(low_conflicts),
                ]

                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                document.add_heading("Trademark Definitions: ", level=2)
                # p = document.add_paragraph(" ")
                # p.paragraph_format.line_spacing = Pt(18)
                p = document.add_paragraph(
                    "CONDITION 1: MARK: NAME-BASED SIMILARITY (comprised of Exact Match, Semantically Equivalent, Phonetically Equivalent, Primary position match)"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 2: CLASS: CLASS OVERLAP")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(
                    "CONDITION 3: GOODS/SERVICES: OVERLAPPING GOODS/SERVICES & TARGET MARKETS"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(
                    "DIRECT HIT: Direct Name hit, regardless of the class"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                # Populate the table with the labels and values
                for i in range(5):
                    table.cell(i, 0).text = labels[i]
                    table.cell(i, 1).text = str(values[i])

                    # Set the font size to 10 for both cells
                    for cell in table.row_cells(i):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 3 conditions satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_high = pd.DataFrame(high_conflicts)
                    df_high = df_high.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_high = document.add_table(
                        df_high.shape[0] + 1, df_high.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_high.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_high.columns):
                        table_high.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_high.iterrows():
                        for j, value in enumerate(row):
                            cell = table_high.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 2 conditions satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_moderate = pd.DataFrame(moderate_conflicts)
                    df_moderate = df_moderate.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_moderate = document.add_table(
                        df_moderate.shape[0] + 1, df_moderate.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_moderate.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_moderate.columns):
                        table_moderate.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_moderate.iterrows():
                        for j, value in enumerate(row):
                            cell = table_moderate.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_Name_Matchs = pd.DataFrame(Name_Matchs)
                    df_Name_Matchs = df_Name_Matchs.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_Name_Matchs = document.add_table(
                        df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_Name_Matchs.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_Name_Matchs.columns):
                        table_Name_Matchs.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_Name_Matchs.iterrows():
                        for j, value in enumerate(row):
                            cell = table_Name_Matchs.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 1 condition satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_low = pd.DataFrame(low_conflicts)
                    df_low = df_low.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                    # Set a predefined table style (with borders)
                    table_low.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_low.columns):
                        table_low.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_low.iterrows():
                        for j, value in enumerate(row):
                            cell = table_low.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                def add_conflict_paragraph(document, conflict):
                    p = document.add_paragraph(
                        f"Trademark Name : {conflict.get('Trademark name', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")
                    p.paragraph_format.line_spacing = Pt(18)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 3 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in high_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 2 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in moderate_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts Reasoning:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in Name_Matchs:
                        add_conflict_paragraph(document, conflict)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 1 condition satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in low_conflicts:
                        add_conflict_paragraph(document, conflict)

                def add_conflict_paragraph_to_array(conflict):
                    result = []
                    result.append(
                        f"Trademark Name : {conflict.get('Trademark name', 'N/A')}"
                    )
                    result.append(
                        f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}"
                    )
                    result.append(
                        f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}"
                    )
                    result.append(
                        f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}"
                    )
                    result.append(
                        f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}"
                    )
                    result.append(
                        f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}"
                    )
                    result.append(
                        f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}"
                    )
                    result.append(" ")  # Blank line for spacing
                    result.append(f"{conflict.get('reasoning', 'N/A')}\n")
                    result.append(" ")  # Blank line for spacing
                    return result

                conflicts_array = []

                if len(high_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 3 conditions satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in high_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(moderate_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 2 conditions satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in moderate_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(Name_Matchs) > 0:
                    conflicts_array.append(
                        "Trademarks with Name Match's Conflicts Reasoning:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in Name_Matchs:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(low_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 1 condition satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in low_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                # for i in range(70,96):
                #     progress_bar.progress(i)

                progress_bar.progress(100)

                filename = proposed_name
                doc_stream = BytesIO()
                document.save(doc_stream)
                doc_stream.seek(0)
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'
                st.sidebar.markdown(download_table, unsafe_allow_html=True)
                st.success(
                    f"{proposed_name} Document conflict report successfully completed!"
                )

                opinion_output = run_trademark_analysis(
                    proposed_name,
                    proposed_class,
                    proposed_goods_services,
                    conflicts_array,
                )
                # Ensure extracted_data is defined by assigning the result of extract_web_common_law
                extracted_data = extract_web_common_law(
                    full_web_common_law, proposed_name
                )
                web_common_law_opinion = analyze_web_common_law(
                    extracted_data, proposed_name
                )
                st.write(
                    "------------------------------------------------------------------------------------------------------------------------------"
                )
                st.write(opinion_output)

                # Export to Word
                filename = export_trademark_opinion_to_word(
                    opinion_output, web_common_law_opinion
                )

                # Download button
                with open(filename, "rb") as file:
                    st.sidebar.download_button(
                        label="Download Trademark Opinion",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60
                st.write(f"Time taken: {elapsed_time} mins")

                st.write(
                    "______________________________________________________________________________________________________________________________"
                )

        progress_bar.progress(100)
        st.success("All documents processed successfully!")
